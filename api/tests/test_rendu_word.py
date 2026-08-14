"""Tests de l'adaptateur Word."""

from __future__ import annotations

import io
from datetime import UTC, datetime

import pytest
from docx import Document as DocxDocument

from app.models.constat import NiveauConfiance
from app.models.rapport import Affirmation, Document, Manifeste, Section, Tableau
from app.services.rendu.word import rendu_word


def _document(mention: str = "", lacune: bool = False) -> Document:
    return Document(
        titre="Étude de filière — le cacao",
        sous_titre="Analyse documentée",
        sections=(
            Section(
                titre="Contexte",
                corps="Un paragraphe analytique.",
                affirmations=(
                    Affirmation(
                        texte="La production avoisine 2,2 millions de tonnes.",
                        source="CNRA",
                        date="2025-10-01",
                        methode="rag",
                        confiance=NiveauConfiance.MOYENNE,
                        empreinte="e3b0c44298fc",
                    ),
                ),
                lacune=lacune,
            ),
        ),
        tableaux=(
            Tableau(titre="Prix", entetes=("Campagne", "Prix"), lignes=(("2025-2026", "1 500"),)),
        ),
        manifeste=Manifeste(
            modele="opencacao-8b",
            version_modele="1.1.0",
            version_app="0.6.75",
            profil_materiel="cpu",
            genere_le=datetime(2026, 7, 29, 10, 0, tzinfo=UTC),
            empreinte_demandeur="a1b2c3d4e5f6",
        ),
        mention=mention,
    )


def _textes(octets: bytes) -> list[str]:
    docx = DocxDocument(io.BytesIO(octets))
    return [paragraphe.text for paragraphe in docx.paragraphs]


def test_le_rendu_est_un_docx_ouvrable():
    """Un docx invalide se voit a la premiere demonstration, pas avant."""
    octets = rendu_word(_document())
    assert octets[:2] == b"PK"  # conteneur ZIP OOXML
    assert _textes(octets)


def test_le_titre_figure_dans_le_document():
    assert "Étude de filière — le cacao" in _textes(rendu_word(_document()))


def test_le_titre_de_section_et_son_corps_figurent():
    textes = _textes(rendu_word(_document()))
    assert "Contexte" in textes
    assert "Un paragraphe analytique." in textes


def test_la_mention_precede_le_contenu():
    """D5 : en tete, pas en annexe."""
    textes = _textes(rendu_word(_document(mention="Document préparatoire.")))
    assert textes.index("Document préparatoire.") < textes.index("Contexte")


def test_une_section_en_lacune_est_signalee():
    assert any("lacune" in texte.lower() for texte in _textes(rendu_word(_document(lacune=True))))


def test_les_tableaux_sont_de_vrais_tableaux_word():
    """Un tableau aligne a l espace n est pas exploitable par le destinataire."""
    docx = DocxDocument(io.BytesIO(rendu_word(_document())))
    assert len(docx.tables) >= 2  # tableau de données + provenance
    assert docx.tables[0].cell(0, 0).text == "Campagne"
    assert docx.tables[0].cell(1, 1).text == "1 500"


def test_le_tableau_de_provenance_porte_chaque_affirmation():
    docx = DocxDocument(io.BytesIO(rendu_word(_document())))
    provenance = docx.tables[-1]
    assert provenance.cell(0, 0).text == "Section"
    assert provenance.cell(1, 2).text == "CNRA"


def test_le_manifeste_figure_dans_le_document():
    textes = " ".join(_textes(rendu_word(_document())))
    assert "opencacao-8b" in textes
    assert "0.6.75" in textes
    assert "a1b2c3d4e5f6" in textes


def test_un_caractere_de_controle_ne_corrompt_pas_le_fichier():
    """python-docx n echappe rien : un \\x00 venant du modele casserait le XML."""
    document = _document()
    sale = Document(
        titre=document.titre,
        sous_titre=document.sous_titre,
        sections=(Section(titre="Contexte", corps="Texte\x00 avec\x0b controle."),),
        tableaux=(Tableau(titre="T", entetes=("A",), lignes=(("x\x00y",),)),),
        manifeste=document.manifeste,
    )
    octets = rendu_word(sale)
    textes = _textes(octets)
    assert any("Texte" in texte for texte in textes)
    assert "\x00" not in "".join(textes)


def test_le_manifeste_ne_porte_jamais_l_identifiant_du_demandeur():
    """Ce test doit ECHOUER si l identifiant est ecrit en plus de l empreinte."""
    assert "appareil-a" not in " ".join(_textes(rendu_word(_document())))


@pytest.mark.parametrize("hostile", ["\ud800", "\udfff", "\ufffe", "\uffff"])
def test_un_caractere_refuse_par_xml_ne_casse_pas_le_fichier(hostile):
    """Sans purge, l echec remonte en ValueError et le routeur rend « format inconnu »."""
    document = _document()
    sale = Document(
        titre=f"Titre {hostile}",
        sous_titre=document.sous_titre,
        sections=(Section(titre=f"Une {hostile}", corps=f"corps {hostile} suite"),),
        tableaux=(Tableau(titre="T", entetes=("A",), lignes=((f"x{hostile}y",),)),),
        manifeste=document.manifeste,
    )
    assert rendu_word(sale)[:2] == b"PK"


def test_les_metadonnees_du_fichier_sont_celles_du_projet():
    """Sans cela, le livrable part chez le bailleur signe « python-docx »."""
    proprietes = DocxDocument(io.BytesIO(rendu_word(_document()))).core_properties
    assert "OpenCacao" in proprietes.author
    assert proprietes.title == "Étude de filière — le cacao"
    assert "python-docx" not in (proprietes.comments or "")
