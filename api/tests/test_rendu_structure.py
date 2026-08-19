"""Un livrable doit s'ouvrir comme un document, pas comme une suite de paragraphes.

Demande de Waopron le 19/08/2026, après lecture d'une étude réellement produite : le
document commençait directement par la première section. Il manquait ce qui fait
qu'un lecteur institutionnel reconnaît une étude — page de garde, sommaire, résumé,
conclusion.

Ces quatre éléments sont **déterministes** : ils viennent de la structure du document,
pas d'une génération. Le résumé et la conclusion, eux, sont rédigés à partir des
sections DÉJÀ écrites et sourcées : c'est une synthèse, jamais un fait nouveau — la
règle « n'invente rien » vaut aussi pour la dernière page.
"""

from __future__ import annotations

import io
from datetime import UTC, datetime

import pytest
from docx import Document as DocxDocument

from app.models.domain import Confiance
from app.models.rapport import Affirmation, Document, Manifeste, Section
from app.services.rendu.markdown import rendu_markdown
from app.services.rendu.word import rendu_word


@pytest.fixture
def document() -> Document:
    champs = {n: "" for n in Manifeste.__dataclass_fields__}
    champs["genere_le"] = datetime(2026, 8, 19, 7, 0, tzinfo=UTC)
    champs["outils"] = ()
    champs["documents_rag"] = ()
    manifeste = Manifeste(**champs)
    return Document(
        titre="Étude de marché — cacao ivoirien",
        sous_titre="Opportunités pour un nouvel acheteur",
        sections=(
            Section(
                titre="Structure du marché",
                corps="Le marché ivoirien est encadré par le Conseil du Café-Cacao.",
                affirmations=(
                    Affirmation(
                        texte="426 acheteurs agréés.",
                        source="Conseil du Café-Cacao",
                        date="",
                        methode="rag",
                        confiance=Confiance.ELEVEE,
                    ),
                ),
            ),
            Section(titre="Prix officiel", corps="Le prix bord champ est fixé chaque campagne."),
        ),
        tableaux=(),
        manifeste=manifeste,
        resume="Le marché ivoirien est régulé et concentré autour d'acheteurs agréés.",
        conclusion="L'accès au marché passe par un acheteur agréé et une coopérative.",
    )


def _texte_word(octets: bytes) -> str:
    docx = DocxDocument(io.BytesIO(octets))
    return "\n".join(p.text for p in docx.paragraphs)


def test_le_document_porte_un_resume_et_une_conclusion(document: Document) -> None:
    """Le modèle de données doit les transporter : sans champ, aucun rendu ne peut."""
    assert document.resume
    assert document.conclusion


@pytest.mark.parametrize("rendu", [rendu_markdown, rendu_word])
def test_le_sommaire_liste_toutes_les_sections(rendu) -> None:
    """Un sommaire incomplet est pire que pas de sommaire : il ment sur le contenu."""
    champs = {n: "" for n in Manifeste.__dataclass_fields__}
    champs["genere_le"] = datetime(2026, 8, 19, 7, 0, tzinfo=UTC)
    champs["outils"] = ()
    champs["documents_rag"] = ()
    doc = Document(
        titre="T",
        sous_titre="S",
        sections=(
            Section(titre="Première section", corps="a"),
            Section(titre="Deuxième section", corps="b"),
            Section(titre="Troisième section", corps="c"),
        ),
        tableaux=(),
        manifeste=Manifeste(**champs),
        resume="R",
        conclusion="C",
    )

    sortie = rendu(doc)
    texte = sortie if isinstance(sortie, str) else _texte_word(sortie)

    assert "Sommaire" in texte
    for titre in ("Première section", "Deuxième section", "Troisième section"):
        assert texte.count(titre) >= 2, f"{titre} doit figurer au sommaire ET dans le corps"


def test_word_place_le_resume_avant_les_sections(document: Document) -> None:
    texte = _texte_word(rendu_word(document))

    assert texte.index("Résumé") < texte.index("Structure du marché")


def test_word_place_la_conclusion_apres_les_sections(document: Document) -> None:
    texte = _texte_word(rendu_word(document))

    assert texte.index("Conclusion") > texte.index("Prix officiel")


def test_word_ouvre_sur_une_page_de_garde(document: Document) -> None:
    """La page de garde doit être SEULE sur sa page : sans saut, le sommaire remonte
    à côté du titre et le document ne ressemble plus à une étude."""
    docx = DocxDocument(io.BytesIO(rendu_word(document)))
    xml = docx.element.xml

    assert 'w:type="page"' in xml, "aucun saut de page après la page de garde"


@pytest.mark.parametrize("rendu", [rendu_markdown, rendu_word])
def test_un_document_sans_resume_ni_conclusion_reste_valide(rendu) -> None:
    """Compatibilité : les livrables courts (bulletin régional) n'en ont pas besoin,
    et ne doivent pas afficher de rubrique vide."""
    champs = {n: "" for n in Manifeste.__dataclass_fields__}
    champs["genere_le"] = datetime(2026, 8, 19, 7, 0, tzinfo=UTC)
    champs["outils"] = ()
    champs["documents_rag"] = ()
    doc = Document(
        titre="Bulletin",
        sous_titre="",
        sections=(Section(titre="Météo", corps="Pluies attendues."),),
        tableaux=(),
        manifeste=Manifeste(**champs),
    )

    sortie = rendu(doc)
    texte = sortie if isinstance(sortie, str) else _texte_word(sortie)

    assert "Résumé" not in texte
    assert "Conclusion" not in texte
