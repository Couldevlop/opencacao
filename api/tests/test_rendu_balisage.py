"""Aucun astérisque de balisage ne doit atteindre un fichier livré.

Audit du 19/08 : 52 occurrences de ``**`` imprimées telles quelles dans le Word.
Ces tests assèrent sur le CONTENU DU FICHIER — ce qu'un lecteur verra — et non sur
l'appel d'une fonction de conversion.
"""

from __future__ import annotations

import io
import re
import zipfile
from datetime import UTC, datetime

import pytest

from app.models.constat import NiveauConfiance
from app.models.rapport import Affirmation, Document, Manifeste, Section, Tableau
from app.services.rendu.diapositives import rendu_pptx
from app.services.rendu.tableur import rendu_excel
from app.services.rendu.word import rendu_word

CORPS = (
    "Le prix officiel s'élève à **1 200 FCFA par kilogramme** pour la campagne, "
    "fixé par le *Conseil du Café-Cacao* avant son ouverture."
)


def _document() -> Document:
    return Document(
        titre="Étude de filière — cacao",
        sous_titre="Recette du balisage",
        sections=(
            Section(
                "Prix officiel en vigueur",
                CORPS,
                (
                    Affirmation(
                        texte="Le prix est **1 200 FCFA/kg**.",
                        source="Conseil du Café-Cacao",
                        date="",
                        methode="outil:prix",
                        confiance=NiveauConfiance.ELEVEE,
                    ),
                ),
            ),
        ),
        tableaux=(Tableau("Chiffres", ("Poste", "Valeur"), (("Prix", "**1 200** FCFA"),)),),
        manifeste=Manifeste(
            modele="opencacao-8b",
            version_modele="1.1.0",
            version_app="0.6.91",
            profil_materiel="gpu",
            genere_le=datetime(2026, 8, 20, 9, 0, tzinfo=UTC),
            empreinte_demandeur="abc",
        ),
        resume="Le résumé cite lui aussi **1 200 FCFA**.",
        conclusion="La conclusion mentionne **50 %** de transformation locale.",
    )


def _textes(octets: bytes, chemin_regex: str) -> str:
    paquet = zipfile.ZipFile(io.BytesIO(octets))
    return " ".join(
        paquet.read(n).decode("utf-8", "replace")
        for n in paquet.namelist()
        if re.search(chemin_regex, n)
    )


def test_aucun_asterisque_de_gras_dans_le_word() -> None:
    """Le défaut mesuré dans les fichiers livrés le 19/08."""
    xml = _textes(rendu_word(_document()), r"^word/document\.xml$")
    visibles = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))
    assert "**" not in visibles
    assert "1 200 FCFA par kilogramme" in visibles


def test_le_gras_est_reellement_gras_dans_le_word() -> None:
    """Retirer les astérisques sans porter le gras serait perdre l'intention.

    On relit le document avec python-docx et on vérifie LE segment concerné : se
    contenter de chercher un `<w:b/>` quelque part passerait déjà sans correctif,
    les en-têtes de tableau étant gras de longue date.
    """
    from docx import Document as DocxDocument

    docx = DocxDocument(io.BytesIO(rendu_word(_document())))
    runs = [r for p in docx.paragraphs for r in p.runs]
    cible = [r for r in runs if r.text == "1 200 FCFA par kilogramme"]
    assert cible, "le segment mis en gras par le modèle doit exister tel quel"
    assert cible[0].bold is True
    voisins = [r for r in runs if r.text.startswith("Le prix officiel")]
    assert voisins and not voisins[0].bold, "le reste de la phrase ne doit PAS être gras"


def test_aucun_asterisque_dans_le_powerpoint() -> None:
    xml = _textes(rendu_pptx(_document()), r"^ppt/slides/slide\d+\.xml$")
    visibles = "".join(re.findall(r"<a:t>([^<]*)</a:t>", xml))
    assert "**" not in visibles


def test_aucun_asterisque_dans_le_tableur() -> None:
    """Excel ne porte pas d'enrichissement : le texte doit y arriver nu."""
    contenu = _textes(rendu_excel(_document()), r"^xl/(sharedStrings|worksheets/).*\.xml$")
    assert "**" not in contenu


@pytest.mark.parametrize("rendu", [rendu_word, rendu_pptx])
def test_le_resume_et_la_conclusion_sont_traites_aussi(rendu) -> None:
    """Ils sont rédigés par le même modèle, avec le même balisage."""
    paquet = zipfile.ZipFile(io.BytesIO(rendu(_document())))
    tout = " ".join(
        paquet.read(n).decode("utf-8", "replace") for n in paquet.namelist() if n.endswith(".xml")
    )
    balises = re.findall(r"<(?:w:t|a:t)[^>]*>([^<]*)</(?:w:t|a:t)>", tout)
    assert "**" not in "".join(balises)
