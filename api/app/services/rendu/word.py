"""Rendu Word — dossier de parcelle et étude de filière.

Les conventions typographiques (tailles, couleurs, hiérarchie de titres) reprennent
``scripts/build_doc_agentique.py`` : le projet a déjà une identité de document écrite,
on ne lui en invente pas une seconde.

**Le texte est purgé de ce que XML 1.0 refuse.** ``python-docx`` n'échappe rien : un
octet nul ou un surrogate isolé venant du modèle produirait un XML invalide, donc un
fichier que le destinataire ne peut pas ouvrir. La contrainte étant commune aux trois
formats OOXML, elle vit dans ``rendu/ooxml.py``.
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.application.provenance import tableau_de_provenance
from app.models.rapport import Document, Graphique, Tableau
from app.services.rendu import charte, enrichi
from app.services.rendu.ooxml import texte_xml_sur
from app.services.rendu.ooxml_graphique import TYPE_CONTENU as TYPE_CONTENU_GRAPHIQUE
from app.services.rendu.ooxml_graphique import partie_graphique

# Auteur inscrit dans les métadonnées du fichier livré.
_AUTEUR = "OpenCacao — OpenLab Consulting"

# Palette reprise de scripts/build_doc_agentique.py.
_ORANGE = RGBColor(0xEA, 0x5B, 0x13)
_SOMBRE = RGBColor(0x1F, 0x1F, 0x1F)
_GRIS = RGBColor(0x60, 0x60, 0x60)
_BLANC = RGBColor(0xFF, 0xFF, 0xFF)

# Aplats. Un document sans couleur se lit comme un mémo ; trop coloré, il perd son
# sérieux. Trois teintes suffisent : l'orange du projet pour ce qui structure, son
# dégradé très clair pour le zébrage, un gris pâle pour ce qui n'est qu'informatif.
_FOND_ORANGE = charte.ORANGE
_FOND_ORANGE_PALE = charte.FOND_ORANGE_PALE
_FOND_GRIS_PALE = charte.FOND_GRIS_PALE

_MOIS = (
    "janvier",
    "février",
    "mars",
    "avril",
    "mai",
    "juin",
    "juillet",
    "août",
    "septembre",
    "octobre",
    "novembre",
    "décembre",
)


def _date_lisible(moment) -> str:
    """Rend une date en toutes lettres, sans dépendre d'une locale système.

    Une page de garde portant « 2026-08-20T09:04:29.675716+00:00 » annonce une machine,
    pas une étude (audit du 19/08). L'horodatage précis reste au manifeste, où il sert
    la rejouabilité.
    """
    return f"{moment.day} {_MOIS[moment.month - 1]} {moment.year}"


def _ombrer(element, couleur: str) -> None:
    """Pose un aplat de couleur sur une cellule ou un paragraphe.

    python-docx n'expose pas l'ombrage : on ajoute le nœud ``w:shd`` aux propriétés,
    ce que Word lit indifféremment sur une cellule (``tcPr``) ou un paragraphe (``pPr``).
    """
    ombre = OxmlElement("w:shd")
    ombre.set(qn("w:val"), "clear")
    ombre.set(qn("w:color"), "auto")
    ombre.set(qn("w:fill"), couleur)
    element.append(ombre)


def _encadre(docx, texte: str, *, fond: str, couleur: RGBColor, gras: bool = False) -> None:
    """Ajoute un paragraphe sur aplat — ce qui doit sauter aux yeux.

    Sert la mention réglementaire (D5) et l'aveu de lacune : deux endroits où le
    document dit quelque chose d'important sur lui-même, et où se fondre dans le corps
    reviendrait à l'enterrer.
    """
    paragraphe = _paragraphe(docx, texte, taille=10, gras=gras, couleur=couleur)
    _ombrer(paragraphe.paragraph_format.element.get_or_add_pPr(), fond)
    paragraphe.paragraph_format.space_before = Pt(6)
    paragraphe.paragraph_format.space_after = Pt(10)


def _propre(texte: str) -> str:
    """Retire ce que XML 1.0 refuse (contrôles, surrogates isolés, non-caractères)."""
    return texte_xml_sur(texte)


def _paragraphe(
    docx,
    texte: str,
    *,
    taille: int = 11,
    gras: bool = False,
    couleur: RGBColor = _SOMBRE,
    italique: bool = False,
):
    """Ajoute un paragraphe à la mise en forme du projet.

    Args:
        docx: Document python-docx en cours de construction.
        texte: Contenu du paragraphe.
        taille: Corps de la police, en points.
        gras: Met le texte en gras.
        couleur: Couleur du texte.
        italique: Met le texte en italique.

    Returns:
        Le paragraphe ajouté.
    """
    paragraphe = docx.add_paragraph()
    # Le balisage du modèle est RENDU, pas imprimé : l'audit du 19/08 comptait 52
    # occurrences de « ** » écrites en toutes lettres dans les documents livrés.
    for fragment, gras_local, italique_local in enrichi.segments(_propre(texte)):
        run = paragraphe.add_run(fragment)
        run.font.size = Pt(taille)
        run.font.bold = gras or gras_local
        run.font.italic = italique or italique_local
        run.font.color.rgb = couleur
    paragraphe.paragraph_format.space_after = Pt(6)
    return paragraphe


def _tableau_word(docx, tableau: Tableau) -> None:
    """Ajoute un tableau natif Word — et non une image ou du texte aligné.

    Args:
        docx: Document python-docx en cours de construction.
        tableau: Tableau à rendre.
    """
    _paragraphe(docx, tableau.titre, taille=12, gras=True, couleur=_ORANGE)
    table = docx.add_table(rows=1, cols=len(tableau.entetes))
    table.style = "Table Grid"
    # En-tête sur aplat orange, texte blanc : sans lui, un tableau Word se lit comme
    # une grille de tableur et l'œil n'y trouve aucune entrée.
    for colonne, entete in enumerate(tableau.entetes):
        cellule = table.cell(0, colonne)
        cellule.text = enrichi.sans_balisage(_propre(entete))
        _ombrer(cellule._tc.get_or_add_tcPr(), _FOND_ORANGE)
        for paragraphe in cellule.paragraphs:
            for run in paragraphe.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = _BLANC
    for rang, ligne in enumerate(tableau.lignes):
        cellules = table.add_row().cells
        for colonne, valeur in enumerate(ligne):
            cellules[colonne].text = enrichi.sans_balisage(_propre(valeur))
            # Zébrage : sur un tableau long, c'est ce qui empêche de sauter une ligne.
            if rang % 2 == 0:
                _ombrer(cellules[colonne]._tc.get_or_add_tcPr(), _FOND_ORANGE_PALE)
            for paragraphe in cellules[colonne].paragraphs:
                for run in paragraphe.runs:
                    run.font.size = Pt(10)


def _identifier(proprietes, document: Document) -> None:
    """Renseigne les métadonnées du fichier.

    Sans cela, le livrable part chez le bailleur signé « python-docx », avec une
    description « generated by python-docx » et une application « Microsoft Macintosh
    Word ». Ce sont des métadonnées visibles dans les propriétés du fichier.

    Args:
        proprietes: Propriétés du document (``core_properties``).
        document: Document assemblé, dont on reprend le titre et l'horodatage.
    """
    proprietes.author = _AUTEUR
    proprietes.last_modified_by = _AUTEUR
    proprietes.title = _propre(document.titre)
    proprietes.comments = ""
    proprietes.created = document.manifeste.genere_le
    proprietes.modified = document.manifeste.genere_le


# Dimensions d'une figure en unités anglaises (EMU) : 914 400 EMU par pouce. 15,24 cm de
# large tient dans les marges par défaut d'un A4 portrait sans déborder.
_LARGEUR_FIGURE = 5486400
_HAUTEUR_FIGURE = 3200400


def _inserer_graphique(docx: Document, graphique: Graphique, rang: int) -> None:
    """Ajoute une figure NATIVE au document Word.

    ``python-docx`` n'a pas d'API de graphique : on crée la partie nous-mêmes, on la
    relie au document, et on insère le cadre qui la référence. Le type de contenu est
    porté par la partie — le rédacteur de paquet écrit alors l'override qui va bien dans
    ``[Content_Types].xml``, sans quoi Word refuse d'ouvrir le fichier.

    Args:
        docx: Document Word en construction.
        graphique: Figure à tracer.
        rang: Rang de la figure, qui fixe le nom de la partie (chart1, chart2…).
    """
    partie = Part(
        PackURI(f"/word/charts/chart{rang}.xml"),
        TYPE_CONTENU_GRAPHIQUE,
        partie_graphique(graphique),
        docx.part.package,
    )
    identifiant = docx.part.relate_to(partie, RT.CHART)
    cadre = parse_xml(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:drawing '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{_LARGEUR_FIGURE}" cy="{_HAUTEUR_FIGURE}"/>'
        '<wp:docPr id="{rang}" name="Figure {rang}"/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
        '<c:chart xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'r:id="{rid}"/></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>'.format(
            rang=rang, rid=identifiant
        )
    )
    docx.element.body.append(cadre)
    if graphique.note:
        # La note dit d'où viennent les chiffres. Sans elle, une figure est une
        # affirmation de plus — exactement ce que ce projet refuse.
        _paragraphe(docx, graphique.note, taille=9, couleur=_GRIS, italique=True)


def _pied_de_page(docx, titre: str) -> None:
    """Pose un pied de page : titre court à gauche, numéro de page à droite.

    Sans lui, impossible de dire « page 4 » en réunion, et une feuille photocopiée
    seule ne dit plus d'où elle vient (audit du 19/08). Le numéro est un CHAMP Word,
    donc recalculé à l'impression — pas un nombre écrit en dur.
    """
    pied = docx.sections[0].footer
    paragraphe = pied.paragraphs[0] if pied.paragraphs else pied.add_paragraph()
    paragraphe.text = ""
    gauche = paragraphe.add_run(f"{_propre(titre)}    ")
    gauche.font.size = Pt(8)
    gauche.font.color.rgb = _GRIS
    paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Champ PAGE : trois nœuds (début, instruction, fin), la forme qu'attend Word.
    run = paragraphe.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = _GRIS
    debut = OxmlElement("w:fldChar")
    debut.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    for noeud in (debut, instruction, fin):
        run._r.append(noeud)


def rendu_word(document: Document) -> bytes:
    """Rend le document au format Word.

    Args:
        document: Document assemblé par le moteur.

    Returns:
        Les octets du fichier ``.docx``.
    """
    docx = DocxDocument()
    _identifier(docx.core_properties, document)
    _pied_de_page(docx, document.titre)

    # --- Page de garde ---------------------------------------------------------
    # Seule sur sa page : sans le saut, le sommaire remonte à côté du titre et le
    # fichier ne ressemble plus à une étude mais à une suite de paragraphes.
    _paragraphe(docx, document.titre, taille=24, gras=True, couleur=_ORANGE)
    if document.sous_titre:
        _paragraphe(docx, document.sous_titre, taille=13, couleur=_GRIS, italique=True)
    if document.mention:
        # D5 : en tête, avant tout contenu, et visuellement distincte — sur aplat, pas
        # seulement en gras orange, qui se fondait dans le reste de la page de garde.
        _encadre(docx, document.mention, fond=_FOND_ORANGE_PALE, couleur=_ORANGE, gras=True)
    _paragraphe(
        docx, f"Généré le {_date_lisible(document.manifeste.genere_le)}", taille=10, couleur=_GRIS
    )
    docx.add_page_break()

    # --- Sommaire --------------------------------------------------------------
    # Construit depuis la structure, jamais demandé au modèle : un sommaire qui ne
    # correspond pas au contenu est pire que pas de sommaire.
    _paragraphe(docx, "Sommaire", taille=16, gras=True, couleur=_ORANGE)
    entrees = list(document.sections)
    if document.resume:
        _paragraphe(docx, "Résumé", taille=11, couleur=_GRIS)
    for rang, section in enumerate(entrees, start=1):
        _paragraphe(docx, f"{rang}. {section.titre}", taille=11, couleur=_GRIS)
    if document.conclusion:
        _paragraphe(docx, "Conclusion", taille=11, couleur=_GRIS)
    _paragraphe(docx, "Annexe — manifeste et provenance", taille=11, couleur=_GRIS)
    docx.add_page_break()

    if document.resume:
        _paragraphe(docx, "Résumé", taille=15, gras=True, couleur=_ORANGE)
        _paragraphe(docx, document.resume)

    for section in document.sections:
        _paragraphe(docx, section.titre, taille=15, gras=True, couleur=_ORANGE)
        if section.lacune:
            # UN seul aveu. Le document en disait deux, l'un après l'autre : bafouiller
            # à l'endroit précis où l'on reconnaît une faiblesse est le pire moment.
            # « Lacune — » nomme la chose UNE fois, puis le corps l'explique. Le
            # document en disait deux fois la même chose, l'un après l'autre.
            _encadre(docx, f"Lacune — {section.corps}", fond=_FOND_GRIS_PALE, couleur=_GRIS)
        else:
            _paragraphe(docx, section.corps)

    if document.conclusion:
        _paragraphe(docx, "Conclusion", taille=15, gras=True, couleur=_ORANGE)
        _paragraphe(docx, document.conclusion)

    if document.graphiques:
        _paragraphe(docx, "Base probante du document", taille=15, gras=True, couleur=_ORANGE)
        for rang, graphique in enumerate(document.graphiques, start=1):
            _inserer_graphique(docx, graphique, rang)

    for tableau in document.tableaux:
        _tableau_word(docx, tableau)

    manifeste = document.manifeste
    _paragraphe(docx, "Annexe — manifeste de génération", taille=15, gras=True, couleur=_ORANGE)
    for ligne in (
        f"Modèle : {manifeste.modele} (version {manifeste.version_modele})",
        f"Version applicative : {manifeste.version_app}",
        f"Profil matériel : {manifeste.profil_materiel}",
        # À la seconde : la microseconde n'ajoute rien à la rejouabilité.
        f"Généré le : {manifeste.genere_le.replace(microsecond=0).isoformat()}",
        f"Empreinte du demandeur : {manifeste.empreinte_demandeur}",
    ):
        _paragraphe(docx, ligne, taille=10, couleur=_GRIS)

    _tableau_word(docx, tableau_de_provenance(document))

    tampon = io.BytesIO()
    docx.save(tampon)
    return tampon.getvalue()
