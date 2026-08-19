"""Construit le dossier de présentation d'OpenCacao en Word ET en PDF.

Document destiné au responsable de la communication d'OpenLab Consulting, qui
doit défendre le produit devant des investisseurs, des bailleurs internationaux,
le Conseil du Café-Cacao, l'ANADER, le CNRA et des ONG. Ton orienté objectif et
efficacité opérationnelle : chaque chiffre sert un argument, aucun chiffre n'est
estimé.

Principe de construction : un modèle de contenu unique (``contenu_dossier()``,
liste de blocs typés) est rendu par DEUX moteurs indépendants, ``python-docx``
pour le .docx et ``reportlab`` (platypus) pour le .pdf. Les graphiques sont
produits une seule fois en PNG par matplotlib, puis embarqués tels quels dans
les deux documents : les deux fichiers sont donc identiques, sans conversion ni
LibreOffice.

Usage : python scripts/build_dossier_presentation.py
Sorties :
    docs/OpenCacao_Dossier_de_presentation.docx
    docs/OpenCacao_Dossier_de_presentation.pdf
    docs/img_presentation/*.png
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import matplotlib

# Backend non interactif : le script tourne sans serveur graphique.
matplotlib.use("Agg")

import matplotlib.pyplot as plt  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch  # noqa: E402
from reportlab.lib import colors  # noqa: E402
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY  # noqa: E402
from reportlab.lib.pagesizes import A4  # noqa: E402
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # noqa: E402
from reportlab.lib.units import cm, inch  # noqa: E402
from reportlab.platypus import (  # noqa: E402
    HRFlowable,
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# --------------------------------------------------------------------------- #
# Palette de marque et chemins                                                #
# --------------------------------------------------------------------------- #

ORANGE_HEX = "EA5B13"
DARK_HEX = "1F1F1F"
GREY_HEX = "606060"
ORANGE_PALE_HEX = "FDF1E8"
ORANGE_LIGHT_HEX = "F8B98A"
LIGNE_HEX = "DEDEDE"
BLANC_HEX = "FFFFFF"

ORANGE = "#" + ORANGE_HEX
DARK = "#" + DARK_HEX
GREY = "#" + GREY_HEX
ORANGE_PALE = "#" + ORANGE_PALE_HEX
ORANGE_LIGHT = "#" + ORANGE_LIGHT_HEX
NEUTRE = "#D9D9D9"
LIGNE = "#" + LIGNE_HEX

DOC_ORANGE = RGBColor(0xEA, 0x5B, 0x13)
DOC_DARK = RGBColor(0x1F, 0x1F, 0x1F)
DOC_GREY = RGBColor(0x60, 0x60, 0x60)
DOC_BLANC = RGBColor(0xFF, 0xFF, 0xFF)

PDF_ORANGE = colors.HexColor("#" + ORANGE_HEX)
PDF_DARK = colors.HexColor("#" + DARK_HEX)
PDF_GREY = colors.HexColor("#" + GREY_HEX)
PDF_ORANGE_PALE = colors.HexColor("#" + ORANGE_PALE_HEX)
PDF_LIGNE = colors.HexColor("#" + LIGNE_HEX)

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "docs" / "OPENLAB.png"
IMG_DIR = ROOT / "docs" / "img_presentation"
OUT_DOCX = ROOT / "docs" / "OpenCacao_Dossier_de_presentation.docx"
OUT_PDF = ROOT / "docs" / "OpenCacao_Dossier_de_presentation.pdf"

TITRE = "OpenCacao"
SOUS_TITRE = "L'intelligence artificielle souveraine au service de la filière cacao ivoirienne"
ACCROCHE = (
    "Dossier de présentation destiné aux partenaires institutionnels, "
    "aux bailleurs et aux investisseurs"
)
DATE_DOC = "19 août 2026"
PIED = "OpenLab Consulting · OpenCacao · Dossier de présentation · " + DATE_DOC

# --------------------------------------------------------------------------- #
# Données mesurées (aucune valeur estimée ; source : relevés du 19/08/2026)    #
# --------------------------------------------------------------------------- #

# Latence de bout en bout, en secondes, avant et après passage sur GPU.
LATENCE_AVANT, LATENCE_APRES = 38.0, 2.3
# Débit de génération, en tokens par seconde.
DEBIT_CPU, DEBIT_GPU = 15, 123
# Qualité logicielle : tests automatisés et couverture (branches comprises).
NB_TESTS, COUVERTURE = 1655, 99.57
# Continuité de service : repli automatique éprouvé en production.
REPLI_DETECTION, REPLI_TOTAL = 50, 75
REPLI_BASCULE = REPLI_TOTAL - REPLI_DETECTION
# Socle de connaissance.
EXTRAITS, PAIRES_QR, VOLUME_SOURCE = 10269, 10000, "5 Go"
NB_AGENTS, NB_DR, NB_ZONES = 7, 10, 60

ORGANISMES = (
    "Conseil du Café-Cacao",
    "ANADER",
    "CNRA",
    "FIRCA",
    "FAO",
    "ICCO",
)

SOMMAIRE = (
    ("Résumé exécutif", "Ce qu'il faut retenir, en une page."),
    ("Neuf démonstrations exécutées en production",
     "Les scénarios qui font la preuve, un par ligne."),
    ("1. Le besoin auquel OpenCacao répond", "Le conseil existe, sa disponibilité manque."),
    ("2. Ce que fait l'outil : sept agents spécialisés", "Le périmètre fonctionnel, agent par agent."),
    ("3. L'architecture technique, du serveur au producteur",
     "Le schéma des couches, du serveur GPU jusqu'aux écrans."),
    ("4. La solution en trois écrans", "Le chat, Ma parcelle, L'atelier, et la console."),
    ("5. Le socle de connaissance de la filière", "5 Go de sources officielles, 10 269 extraits."),
    ("6. Souveraineté : où vont les données", "Aucun service d'IA externe en production."),
    ("7. Les garde-fous métier", "Ce que l'outil refuse, et vers qui il oriente."),
    ("8. Vision et parcelle : décrire, jamais diagnostiquer", "Photo, contour GPS, exigence EUDR."),
    ("9. Des livrables traçables et rejouables", "Word, Excel, PowerPoint et manifeste de génération."),
    ("10. Performance mesurée", "2,3 secondes de bout en bout, 123 tokens par seconde."),
    ("11. Continuité de service éprouvée", "Un repli automatique de 75 secondes, sans intervention."),
    ("12. Qualité logicielle vérifiable", "1 655 tests, 99,57 % de couverture."),
    ("13. Ce qui reste à établir", "Les chiffres que ce dossier n'avance pas."),
    ("Conclusion", "L'argument à retenir devant un décideur."),
    ("Contact et appel à contribution", "Comment enrichir la base de connaissances."),
)

# Scénarios exécutés en production le 19/08/2026. Formulations conservées telles
# qu'observées : aucun ajout, aucune extrapolation.
ENTETES_SCENARIOS = ["Ce qu'on demande", "Ce que l'outil répond", "Pourquoi c'est important"]
LARGEURS_SCENARIOS = [0.25, 0.43, 0.32]

TABLEAU_SCENARIOS_REFUS = (
    ENTETES_SCENARIOS,
    [
        [
            "« Quel dosage de fongicide dois-je appliquer sur mes cabosses ? »",
            "Refus, avec orientation vers l'agent ANADER local et le Conseil du Café-Cacao. "
            "Aucune dose n'est donnée, même sous insistance, même face à une tentative "
            "d'injection de consigne : le refus tient.",
            "La prescription reste à l'agent habilité. Une dose erronée se paie sur la récolte.",
        ],
        [
            "« Je veux faire une plantation de cacao à Ouangolo. »",
            "L'outil corrige : Ouangolodougou est en savane du Nord, climat trop sec et saison "
            "des pluies trop courte pour le cacaoyer. Il nomme la localité au lieu de dire "
            "« cette zone ». Lorsqu'il ne sait pas si une localité est cacaoyère, il le dit.",
            "Un producteur qui plante au mauvais endroit perd sa mise. Un outil généraliste ne "
            "refuse pas cette erreur.",
        ],
        [
            "« Comment traiter mon champ de maïs contre les chenilles ? »",
            "« Je suis spécialisé UNIQUEMENT dans la filière cacao », avec orientation ANADER. "
            "Même réponse pour l'anacarde.",
            "Le périmètre est tenu : l'outil ne déborde jamais de son domaine de compétence.",
        ],
    ],
    LARGEURS_SCENARIOS,
)

TABLEAU_SCENARIOS_TERRAIN = (
    ENTETES_SCENARIOS,
    [
        [
            "Une photo de plantation déposée depuis un téléphone",
            "En 4,1 secondes : « les cabosses montrent une bonne diversité de couleurs (vert, "
            "rouge, jaune-orangé), signe d'une production active ; les feuilles sont en bon état "
            "et le sol reste couvert », puis « parlez-en à votre agent ANADER local ». Aucun nom "
            "de maladie n'est prononcé.",
            "Choix assumé et verrouillé dans le code : un modèle généraliste est un bon "
            "descripteur et un mauvais diagnosticien. Un producteur qui traite sur un mauvais "
            "diagnostic perd sa récolte.",
        ],
        [
            "Une photo floue",
            "Écartée avant toute analyse, avec un conseil actionnable : « Approchez-vous de la "
            "cabosse, tenez le téléphone bien immobile, et refaites la photo. »",
            "L'outil ne prétend pas lire ce qu'il ne voit pas.",
        ],
        [
            "« Quel est le prix officiel du cacao ? »",
            "1 200 F CFA/kg, campagne intermédiaire 2026, source Conseil du Café-Cacao.",
            "Le prix officiel l'emporte toujours sur les prix historiques présents dans les "
            "documents.",
        ],
        [
            "« Va-t-il pleuvoir cette semaine à Daloa ? »",
            "La prévision réelle : 26,3 mm sur 24 heures. Sans localité précisée, l'outil "
            "demande la commune au lieu de fabriquer une prévision.",
            "Une prévision inventée est un mauvais conseil. L'outil préfère demander.",
        ],
    ],
    LARGEURS_SCENARIOS,
)

TABLEAU_SCENARIOS_SERVICE = (
    ENTETES_SCENARIOS,
    [
        [
            "Le matériel d'inférence coupé volontairement en production",
            "Détection en 50 secondes, retour au service sur le matériel de secours en 75 "
            "secondes au total, sans intervention humaine. L'écran a expliqué la situation aux "
            "utilisateurs pendant l'incident.",
            "La démonstration ne dépend pas d'une machine unique, et l'utilisateur n'est jamais "
            "laissé devant une page muette.",
        ],
        [
            "Une étude produite par l'outil",
            "Un manifeste joint au document : modèle, version, documents mobilisés, outils "
            "appelés, horodatage. Une section sans source disponible est déclarée en lacune.",
            "Le livrable est rejouable et auditable, pas un texte d'origine incertaine.",
        ],
    ],
    LARGEURS_SCENARIOS,
)

TABLEAU_ARCHITECTURE = (
    ["Couche", "Composant", "Ce qu'il fait"],
    [
        ["1. Calcul", "Serveur GPU dédié (vLLM)",
         "Sert le modèle en production : 123 tokens par seconde, réponse en 2,3 secondes"],
        ["1. Calcul", "Repli processeur (llama.cpp, GGUF Q4_K_M)",
         "Reprend le service si le serveur GPU disparaît : 15 tokens par seconde, même modèle"],
        ["1. Calcul", "Ollama",
         "Variante de service local du même export quantifié, pour une démonstration "
         "hors infrastructure"],
        ["2. Modèle", "Ministral 3 8B Instruct et adaptateur LoRA 4 bits",
         "Répond en français simple sur le cacao ivoirien ; poids ouverts, hébergés par "
         "le projet"],
        ["3. Connaissance", "Service d'embeddings (Qwen3-Embedding 0.6B)",
         "Transforme extraits et questions en vecteurs de 1 024 dimensions"],
        ["3. Connaissance", "Index vectoriel souverain",
         "Retrouve les extraits utiles parmi 10 269 ; composant interne, aucun service tiers"],
        ["4. Orchestration", "Garde-fous métier",
         "Filtrent l'entrée et la sortie : les refus non négociables s'appliquent avant tout"],
        ["4. Orchestration", "Orchestrateur d'agents",
         "Route la question vers l'agent compétent, puis assemble la réponse et ses sources"],
        ["4. Orchestration", "Cache et limitation de débit",
         "Servent instantanément les questions récurrentes et protègent le service"],
        ["5. Écrans", "Le chat, Ma parcelle, L'atelier",
         "Trois usages à une seule adresse, dans le navigateur d'un téléphone ou d'un ordinateur"],
        ["5. Écrans", "Console d'administration",
         "Alimente la base de connaissances : recherche des documents puis vectorisation "
         "(accès réservé)"],
    ],
    [0.16, 0.31, 0.53],
)

TABLEAU_MODULES = (
    ["Module", "Ce que l'utilisateur y fait", "Ce que le module produit"],
    [
        [
            "Le chat",
            "Pose sa question en français courant, en une phrase, et poursuit le dialogue",
            "Une réponse fondée sur les documents officiels, avec ses sources et le renvoi "
            "vers l'agent ANADER local",
        ],
        [
            "Ma parcelle",
            "Enregistre une parcelle, fait le tour au GPS depuis son téléphone, dépose "
            "des photos",
            "Un contour, une superficie calculée, des photos datées et un constat satellite "
            "de non-déforestation postérieure au 31/12/2020",
        ],
        [
            "L'atelier",
            "Décrit en une phrase le document attendu, sans remplir de formulaire",
            "Une étude Word, un tableau Excel ou une présentation PowerPoint, accompagnés "
            "du manifeste de génération",
        ],
        [
            "Console d'administration (accès réservé)",
            "Lance la recherche de documents officiels et suit la constitution de la base",
            "Des extraits vectorisés ajoutés à l'index, et un corpus prêt pour le prochain "
            "affinage du modèle",
        ],
    ],
    [0.17, 0.36, 0.47],
)

TABLEAU_VALEUR = (
    ["Besoin", "Bénéficiaire", "Ce qu'apporte OpenCacao", "Élément vérifiable"],
    [
        [
            "Obtenir une réponse fondée sur le corpus officiel",
            "Producteur, technicien",
            "Conseil documentaire sur 10 269 extraits, avec citation des sources",
            "Réponse complète en 2,3 secondes",
        ],
        [
            "Savoir à qui s'adresser près de chez soi",
            "Producteur",
            "Mise en relation avec la Direction Régionale ANADER et la zone",
            "10 directions régionales, 60 zones couvertes",
        ],
        [
            "Disposer du prix officiel sans ambiguïté",
            "Producteur, coopérative",
            "Agent dédié pour lequel le prix officiel prime sur toute valeur historique",
            "Comportement vérifié en production",
        ],
        [
            "Documenter une parcelle et son historique",
            "Coopérative, exportateur",
            "Contour GPS, superficie calculée, photos datées, constat satellite",
            "Non-déforestation postérieure au 31/12/2020 (exigence EUDR)",
        ],
        [
            "Produire une étude ou un dossier présentable",
            "Bailleur, institution, ONG",
            "Livrables Word, Excel et PowerPoint générés avec leur manifeste",
            "Document rejouable à l'identique",
        ],
        [
            "Ne pas exposer les données des producteurs",
            "Filière, autorités de tutelle",
            "Traitement intégral sur l'infrastructure du projet",
            "Aucun service d'IA externe dans le chemin de production",
        ],
    ],
    [0.24, 0.15, 0.35, 0.26],
)

TABLEAU_AGENTS = (
    ["Agent", "Ce qu'il fait", "Sa limite assumée"],
    [
        [
            "Conseil documentaire",
            "Recherche dans la base vectorisée et répond en citant les documents mobilisés",
            "Ne répond pas hors du corpus disponible ; une absence de source est déclarée",
        ],
        [
            "Météo",
            "Rattache la question à une localité cacaoyère et fournit la situation observée",
            "Ne produit jamais une météo qu'il n'a pas reçue",
        ],
        [
            "Prix officiel",
            "Restitue le prix officiel en vigueur",
            "Le prix officiel prime sur toute valeur historique lue dans les documents",
        ],
        [
            "Réglementation",
            "Restitue le cadre applicable à la filière tel qu'il est écrit dans les textes",
            "Ne se substitue pas à un avis juridique",
        ],
        [
            "Normes et certifications",
            "Explique les exigences documentaires des référentiels de la filière",
            "Ne délivre aucune certification, ne préjuge d'aucune décision de certificateur",
        ],
        [
            "Satellite",
            "Signale les alertes de déforestation sur une zone ou une parcelle",
            "Ne produit jamais une attestation de conformité EUDR",
        ],
        [
            "Rédaction de livrables",
            "Compose études et dossiers en Word, Excel et PowerPoint",
            "Une section sans source disponible est déclarée en lacune, jamais estimée",
        ],
    ],
    [0.19, 0.42, 0.39],
)

TABLEAU_GARDE_FOUS = (
    ["Règle", "Ce que l'outil refuse", "Ce qu'il fait à la place"],
    [
        [
            "Aucun dosage phytosanitaire",
            "Toute indication de dose, de concentration ou de cadence de traitement",
            "Oriente vers l'agent ANADER, seul habilité à prescrire sur le terrain",
        ],
        [
            "Le cacao et uniquement le cacao",
            "Toute autre culture, vivrière comme de rente",
            "Explique le périmètre et oriente vers l'ANADER pour les autres cultures",
        ],
        [
            "Aucun avis médical ou vétérinaire",
            "Toute question de santé humaine ou animale",
            "Oriente vers les services de santé compétents",
        ],
        [
            "Aucun diagnostic à partir d'une photo",
            "Toute désignation de maladie sur la base d'une image",
            "Décrit ce qui est visible et renvoie le diagnostic à l'agent ANADER",
        ],
        [
            "Aucune réponse sans point de contact",
            "Une réponse qui laisserait le producteur sans interlocuteur",
            "Joint à chaque réponse le renvoi vers l'agent ANADER local",
        ],
    ],
    [0.22, 0.38, 0.40],
)

TABLEAU_VISION = (
    ["Élément du dossier de parcelle", "Ce que l'outil produit", "Ce qu'il ne produit jamais"],
    [
        [
            "Photo de plantation",
            "Une description objective de ce qui est visible, recoupée avec la météo et "
            "l'historique de la parcelle",
            "Aucun nom de maladie, aucun diagnostic : ils restent à l'agent ANADER",
        ],
        [
            "Contour de la parcelle",
            "Un relevé GPS effectué depuis un téléphone et la superficie calculée",
            "Aucun acte foncier, aucune délimitation opposable à un tiers",
        ],
        [
            "Suivi dans le temps",
            "Des photos datées, rattachées à la parcelle",
            "Aucune valeur estimée : ce qui manque est déclaré en lacune",
        ],
        [
            "Déforestation",
            "Un constat satellite de non-déforestation postérieure au 31/12/2020",
            "Aucune attestation de conformité EUDR : elle relève des organismes habilités",
        ],
    ],
    [0.22, 0.40, 0.38],
)

TABLEAU_MANIFESTE = (
    ["Champ du manifeste", "Ce qu'il garantit au lecteur du livrable"],
    [
        ["Modèle et version", "Savoir exactement quel outil a produit le document"],
        ["Documents mobilisés", "Remonter à la source officielle de chaque affirmation"],
        ["Outils appelés", "Distinguer ce qui vient d'un document de ce qui vient d'une mesure"],
        ["Horodatage", "Situer le livrable dans le temps et le rejouer à l'identique"],
    ],
    [0.30, 0.70],
)

TABLEAU_PERFORMANCE = (
    ["Indicateur", "Avant", "Aujourd'hui", "Lecture"],
    [
        ["Réponse complète de bout en bout", "38 s", "2,3 s", "Le dialogue devient possible"],
        ["Débit de génération", "15 tokens/s", "123 tokens/s", "Plus de huit fois plus rapide"],
        ["Matériel de service", "Processeur", "Carte graphique", "Même modèle, même corpus"],
    ],
    [0.34, 0.16, 0.20, 0.30],
)

TABLEAU_REPLI = (
    ["Séquence", "Durée mesurée", "Ce qui se passe"],
    [
        ["Détection de la défaillance", "50 s", "La supervision constate la perte du service"],
        ["Bascule et retour au service", "25 s", "Le matériel de secours reprend la charge"],
        ["Indisponibilité totale", "75 s", "Aucune intervention humaine dans la boucle"],
    ],
    [0.34, 0.18, 0.48],
)

# --------------------------------------------------------------------------- #
# Graphiques matplotlib                                                       #
# --------------------------------------------------------------------------- #


def _style_matplotlib() -> None:
    """Applique la charte typographique du projet aux figures matplotlib."""
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 10,
            "text.color": DARK,
            "axes.labelcolor": DARK,
            "axes.edgecolor": LIGNE,
            "xtick.color": GREY,
            "ytick.color": GREY,
            "axes.titlecolor": DARK,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "savefig.facecolor": "white",
        }
    )


def _sauver(fig: Any, nom: str) -> Path:
    """Enregistre une figure en PNG haute définition dans docs/img_presentation.

    Args:
        fig: Figure matplotlib à enregistrer.
        nom: Nom de fichier, extension comprise.

    Returns:
        Chemin absolu du fichier écrit.
    """
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    chemin = IMG_DIR / nom
    fig.savefig(chemin, dpi=200, bbox_inches="tight", pad_inches=0.12)
    plt.close(fig)
    return chemin


def figure_chiffres_cles() -> Path:
    """Génère la bande de chiffres clés (six tuiles de synthèse).

    Returns:
        Chemin du PNG produit.
    """
    tuiles = [
        (VOLUME_SOURCE, "de documents de la filière\ntraités pour construire l'outil"),
        (f"{EXTRAITS:,}".replace(",", " "), "extraits vectorisés issus\nde sources officielles"),
        (f"{PAIRES_QR:,}".replace(",", " "), "paires questions/réponses\ncacao pour l'affinage"),
        (str(NB_AGENTS), "agents spécialisés\nen production"),
        (f"{NB_DR} / {NB_ZONES}", "directions régionales ANADER\net zones couvertes"),
        ("0", "service d'IA externe dans\nle chemin de production"),
    ]
    fig = plt.figure(figsize=(9.4, 3.5))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 3)
    ax.set_ylim(0, 2)
    ax.axis("off")
    for index, (valeur, legende) in enumerate(tuiles):
        col, ligne = index % 3, 1 - index // 3
        x, y = col + 0.05, ligne + 0.06
        ax.add_patch(
            FancyBboxPatch(
                (x, y),
                0.9,
                0.88,
                boxstyle="round,pad=0.02,rounding_size=0.05",
                linewidth=1.1,
                edgecolor=ORANGE_LIGHT,
                facecolor=ORANGE_PALE,
            )
        )
        # Filet orange en haut de tuile : rappel de la charte.
        ax.add_patch(
            FancyBboxPatch(
                (x + 0.06, y + 0.79),
                0.16,
                0.035,
                boxstyle="round,pad=0.005,rounding_size=0.01",
                linewidth=0,
                facecolor=ORANGE,
            )
        )
        ax.text(x + 0.45, y + 0.52, valeur, ha="center", va="center", fontsize=27,
                fontweight="bold", color=ORANGE)
        ax.text(x + 0.45, y + 0.22, legende, ha="center", va="center", fontsize=9.5, color=DARK)
    return _sauver(fig, "chiffres_cles.png")


def figure_performance() -> Path:
    """Génère les barres comparatives de latence et de débit.

    Returns:
        Chemin du PNG produit.
    """
    fig, (gauche, droite) = plt.subplots(1, 2, figsize=(9.2, 3.9))

    barres = gauche.bar(
        ["Avant\n(processeur)", "Aujourd'hui\n(carte graphique)"],
        [LATENCE_AVANT, LATENCE_APRES],
        color=[NEUTRE, ORANGE],
        width=0.55,
    )
    gauche.set_title("Réponse complète de bout en bout", fontsize=11, fontweight="bold", pad=12)
    gauche.set_ylabel("secondes")
    gauche.set_ylim(0, LATENCE_AVANT * 1.28)
    for barre, valeur in zip(barres, [LATENCE_AVANT, LATENCE_APRES]):
        etiquette = f"{valeur:.0f} s" if valeur >= 10 else f"{valeur:.1f} s".replace(".", ",")
        gauche.text(barre.get_x() + barre.get_width() / 2, valeur + LATENCE_AVANT * 0.03,
                    etiquette, ha="center", fontsize=13, fontweight="bold", color=DARK)
    gauche.text(0.5, LATENCE_AVANT * 1.14, f"divisé par {LATENCE_AVANT / LATENCE_APRES:.1f}".replace(".", ","),
                ha="center", fontsize=10.5, color=ORANGE, fontweight="bold")

    barres = droite.bar(
        ["Processeur", "Carte graphique"],
        [DEBIT_CPU, DEBIT_GPU],
        color=[NEUTRE, ORANGE],
        width=0.55,
    )
    droite.set_title("Débit de génération", fontsize=11, fontweight="bold", pad=12)
    droite.set_ylabel("tokens par seconde")
    droite.set_ylim(0, DEBIT_GPU * 1.28)
    for barre, valeur in zip(barres, [DEBIT_CPU, DEBIT_GPU]):
        droite.text(barre.get_x() + barre.get_width() / 2, valeur + DEBIT_GPU * 0.03,
                    str(valeur), ha="center", fontsize=13, fontweight="bold", color=DARK)
    droite.text(0.5, DEBIT_GPU * 1.14, f"multiplié par {DEBIT_GPU / DEBIT_CPU:.1f}".replace(".", ","),
                ha="center", fontsize=10.5, color=ORANGE, fontweight="bold")

    for axe in (gauche, droite):
        axe.spines[["top", "right"]].set_visible(False)
        axe.grid(axis="y", color=LIGNE, linewidth=0.7)
        axe.set_axisbelow(True)
    fig.suptitle("Performance mesurée le 19/08/2026", fontsize=12.5, fontweight="bold", y=1.03)
    fig.tight_layout()
    return _sauver(fig, "performance.png")


def figure_couverture() -> Path:
    """Génère le camembert en anneau de la couverture de tests.

    Returns:
        Chemin du PNG produit.
    """
    fig, ax = plt.subplots(figsize=(5.0, 4.2))
    ax.pie(
        [COUVERTURE, 100 - COUVERTURE],
        colors=[ORANGE, "#EFEFEF"],
        startangle=90,
        counterclock=False,
        wedgeprops={"width": 0.32, "edgecolor": "white", "linewidth": 2},
    )
    ax.text(0, 0.12, "99,57 %", ha="center", va="center", fontsize=27,
            fontweight="bold", color=ORANGE)
    ax.text(0, -0.16, "de couverture\nbranches comprises", ha="center", va="center",
            fontsize=10, color=DARK)
    ax.text(0, -1.32, f"{NB_TESTS:,}".replace(",", " ") + " tests automatisés",
            ha="center", fontsize=11.5, fontweight="bold", color=DARK)
    return _sauver(fig, "couverture_tests.png")


def figure_repli() -> Path:
    """Génère le camembert et la chronologie du repli automatique.

    Returns:
        Chemin du PNG produit.
    """
    fig, (gauche, droite) = plt.subplots(
        1, 2, figsize=(9.2, 3.6), gridspec_kw={"width_ratios": [1, 1.5]}
    )

    gauche.pie(
        [REPLI_DETECTION, REPLI_BASCULE],
        labels=[f"Détection\n{REPLI_DETECTION} s", f"Bascule et retour\nau service\n{REPLI_BASCULE} s"],
        colors=[ORANGE, ORANGE_LIGHT],
        startangle=90,
        counterclock=False,
        autopct=lambda part: f"{part:.0f} %",
        textprops={"fontsize": 9.5, "color": DARK},
        wedgeprops={"edgecolor": "white", "linewidth": 2},
    )
    gauche.set_title("Répartition des 75 secondes", fontsize=11, fontweight="bold", pad=14)

    droite.broken_barh(
        [(0, REPLI_DETECTION), (REPLI_DETECTION, REPLI_BASCULE)],
        (0.35, 0.3),
        facecolors=[ORANGE, ORANGE_LIGHT],
        edgecolor="white",
        linewidth=1.5,
    )
    droite.set_xlim(0, REPLI_TOTAL * 1.06)
    droite.set_ylim(0, 1)
    droite.set_yticks([])
    droite.set_xticks([0, 25, REPLI_DETECTION, REPLI_TOTAL])
    droite.set_xticklabels(["0 s", "25 s", "50 s", "75 s"])
    droite.set_xlabel("temps écoulé depuis la défaillance")
    droite.spines[["top", "right", "left"]].set_visible(False)
    droite.text(REPLI_DETECTION / 2, 0.5, "détection", ha="center", va="center",
                fontsize=10, color="white", fontweight="bold")
    droite.text(REPLI_DETECTION + REPLI_BASCULE / 2, 0.5, "bascule", ha="center", va="center",
                fontsize=9.5, color=DARK, fontweight="bold")
    droite.annotate(
        "service rétabli",
        xy=(REPLI_TOTAL, 0.35),
        xytext=(REPLI_TOTAL * 0.72, 0.13),
        fontsize=10,
        color=ORANGE,
        fontweight="bold",
        arrowprops={"arrowstyle": "->", "color": ORANGE, "linewidth": 1.2},
    )
    droite.text(0, 0.82, "Aucune intervention humaine dans la boucle",
                fontsize=10, color=GREY, style="italic")
    droite.set_title("Chronologie du repli, éprouvée en production",
                     fontsize=11, fontweight="bold", pad=14)
    fig.tight_layout()
    return _sauver(fig, "repli_automatique.png")


def figure_chaine() -> Path:
    """Génère le schéma du chemin d'une question, des garde-fous à la réponse.

    Returns:
        Chemin du PNG produit.
    """
    fig = plt.figure(figsize=(9.4, 3.1))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 10)
    ax.set_ylim(0.35, 3.2)
    ax.axis("off")

    etapes = [
        (0.15, "Question du\nproducteur", ORANGE_PALE, DARK),
        (2.55, "Garde-fous\nmétier", ORANGE, "white"),
        (4.95, "Routage vers\nl'agent utile", ORANGE_PALE, DARK),
    ]
    for x, texte, fond, encre in etapes:
        ax.add_patch(
            FancyBboxPatch(
                (x, 1.9),
                1.9,
                0.85,
                boxstyle="round,pad=0.04,rounding_size=0.12",
                linewidth=1.2,
                edgecolor=ORANGE,
                facecolor=fond,
            )
        )
        ax.text(x + 0.95, 2.33, texte, ha="center", va="center", fontsize=10.5,
                color=encre, fontweight="bold")
    ax.add_patch(
        FancyBboxPatch(
            (7.35, 1.9),
            2.45,
            0.85,
            boxstyle="round,pad=0.04,rounding_size=0.12",
            linewidth=1.4,
            edgecolor=DARK,
            facecolor=DARK,
        )
    )
    ax.text(8.57, 2.33, "Réponse sourcée\n+ renvoi ANADER local", ha="center", va="center",
            fontsize=10.5, color="white", fontweight="bold")

    for depart, arrivee in ((2.10, 2.50), (4.50, 4.90), (6.90, 7.30)):
        ax.add_patch(
            FancyArrowPatch(
                (depart, 2.33),
                (arrivee, 2.33),
                arrowstyle="-|>",
                mutation_scale=14,
                linewidth=1.4,
                color=GREY,
            )
        )

    agents = [
        "Conseil documentaire",
        "Météo",
        "Prix officiel",
        "Réglementation",
        "Normes et certifications",
        "Satellite",
        "Rédaction de livrables",
    ]
    # Largeurs proportionnelles à la longueur du libellé, normalisées pour que la
    # rangée tienne exactement dans la largeur du schéma.
    ecart = 0.16
    disponible = 9.7 - ecart * (len(agents) - 1)
    poids = [len(nom) + 3 for nom in agents]
    largeurs = [disponible * p / sum(poids) for p in poids]
    x = 0.15
    for nom, largeur in zip(agents, largeurs):
        ax.add_patch(
            FancyBboxPatch(
                (x, 0.55),
                largeur,
                0.5,
                boxstyle="round,pad=0.03,rounding_size=0.14",
                linewidth=1,
                edgecolor=ORANGE_LIGHT,
                facecolor="white",
            )
        )
        ax.text(x + largeur / 2, 0.80, nom, ha="center", va="center", fontsize=8.4, color=DARK)
        x += largeur + ecart
    ax.text(0.15, 1.42, "Sept agents spécialisés, tous servis par le même modèle souverain",
            fontsize=9.6, color=GREY, style="italic")
    ax.text(0.15, 3.02, "Le chemin d'une question", fontsize=12.5, fontweight="bold", color=DARK)
    return _sauver(fig, "chaine_traitement.png")


def _couche(ax: Any, y: float, hauteur: float, titre: str, chips: list[str],
            *, fond: str = "white", bord: str = ORANGE, x: float = 0.2,
            largeur: float = 6.5, taille_chip: float = 8.2) -> None:
    """Dessine une couche d'architecture et les composants qu'elle contient.

    Args:
        ax: Axes matplotlib cible.
        y: Ordonnée du bas de la couche.
        hauteur: Hauteur de la couche.
        titre: Intitulé de la couche.
        chips: Libellés des composants alignés dans la couche.
        fond: Couleur de fond de la couche.
        bord: Couleur du liseré.
        x: Abscisse du bord gauche.
        largeur: Largeur de la couche.
        taille_chip: Corps du texte des composants.
    """
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            largeur,
            hauteur,
            boxstyle="round,pad=0.03,rounding_size=0.10",
            linewidth=1.2,
            edgecolor=bord,
            facecolor=fond,
        )
    )
    ax.text(x + 0.14, y + hauteur - 0.20, titre, fontsize=8.6, fontweight="bold", color=ORANGE)
    if not chips:
        return
    ecart = 0.12
    interieur = largeur - 0.28 - ecart * (len(chips) - 1)
    poids = [len(c) + 4 for c in chips]
    largeurs = [interieur * p / sum(poids) for p in poids]
    cx = x + 0.14
    for libelle, largeur_chip in zip(chips, largeurs):
        ax.add_patch(
            FancyBboxPatch(
                (cx, y + 0.16),
                largeur_chip,
                hauteur - 0.55,
                boxstyle="round,pad=0.02,rounding_size=0.08",
                linewidth=0.9,
                edgecolor=ORANGE_LIGHT,
                facecolor=ORANGE_PALE,
            )
        )
        ax.text(cx + largeur_chip / 2, y + 0.16 + (hauteur - 0.55) / 2, libelle,
                ha="center", va="center", fontsize=taille_chip, color=DARK)
        cx += largeur_chip + ecart


def figure_architecture() -> Path:
    """Génère le schéma d'architecture technique, du serveur au producteur.

    Returns:
        Chemin du PNG produit.
    """
    fig = plt.figure(figsize=(9.4, 5.2))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")

    # Couches de service, de la fondation matérielle jusqu'aux écrans.
    _couche(ax, 0.25, 1.05, "1. Calcul et service du modèle",
            ["Serveur GPU dédié\nvLLM", "Repli CPU\nllama.cpp (GGUF Q4_K_M)",
             "Variante locale\nOllama"])
    _couche(ax, 1.55, 1.00, "2. Modèle souverain",
            ["Ministral 3 8B Instruct", "Adaptateur LoRA 4 bits (cacao ivoirien)"])
    _couche(ax, 2.80, 1.05, "3. Connaissance vectorisée",
            ["Service d'embeddings\nQwen3-Embedding 0.6B",
             "Index vectoriel souverain\n10 269 extraits, 1 024 dimensions"])
    _couche(ax, 4.10, 1.05, "4. API et orchestration (FastAPI)",
            ["Garde-fous métier", "Orchestrateur\n7 agents", "Cache et\nlimitation de débit"])
    _couche(ax, 5.40, 0.95, "5. Écrans, sur téléphone ou ordinateur",
            ["Le chat", "Ma parcelle", "L'atelier", "Console d'administration\n(accès réservé)"])

    for y_depart in (1.32, 2.58, 3.88, 5.18):
        ax.add_patch(
            FancyArrowPatch(
                (3.45, y_depart),
                (3.45, y_depart + 0.22),
                arrowstyle="-|>",
                mutation_scale=12,
                linewidth=1.3,
                color=GREY,
            )
        )

    # Chaîne documentaire : elle alimente la couche de connaissance.
    ax.add_patch(
        FancyBboxPatch(
            (7.05, 0.25),
            2.75,
            6.10,
            boxstyle="round,pad=0.03,rounding_size=0.10",
            linewidth=1.2,
            edgecolor=DARK,
            facecolor="white",
        )
    )
    ax.text(7.19, 6.20, "Chaîne documentaire", fontsize=8.6, fontweight="bold", color=DARK)
    etapes = [
        "Sources officielles\n(liste blanche de domaines)",
        "Découverte automatique\ndes documents publiés",
        "Téléchargement\nvers le magasin de documents",
        "Extraction et découpage\nen extraits",
        "OCR des documents scannés\n(prévu, hors ligne)",
        "Vectorisation et ajout\nà l'index (jamais de retrait)",
    ]
    y = 5.45
    for index, etape in enumerate(etapes):
        prevu = "prévu" in etape
        ax.add_patch(
            FancyBboxPatch(
                (7.19, y),
                2.47,
                0.62,
                boxstyle="round,pad=0.02,rounding_size=0.08",
                linewidth=0.9,
                linestyle="dashed" if prevu else "solid",
                edgecolor=GREY if prevu else ORANGE_LIGHT,
                facecolor="white" if prevu else ORANGE_PALE,
            )
        )
        ax.text(8.42, y + 0.31, etape, ha="center", va="center", fontsize=7.8,
                color=GREY if prevu else DARK, style="italic" if prevu else "normal")
        if index < len(etapes) - 1:
            ax.add_patch(
                FancyArrowPatch(
                    (8.42, y - 0.02),
                    (8.42, y - 0.24),
                    arrowstyle="-|>",
                    mutation_scale=10,
                    linewidth=1.1,
                    color=GREY,
                )
            )
        y -= 0.86
    ax.add_patch(
        FancyArrowPatch(
            (7.05, 3.32),
            (6.75, 3.32),
            arrowstyle="-|>",
            mutation_scale=13,
            linewidth=1.4,
            color=ORANGE,
        )
    )
    ax.text(0.2, 0.02, "Aucun service d'IA externe n'intervient dans ce schéma : "
                       "chaque couche est opérée par le projet.",
            fontsize=8.6, color=GREY, style="italic")
    return _sauver(fig, "architecture_technique.png")


def _ecran(ax: Any, x: float, titre: str, lignes: list[tuple[str, str]]) -> None:
    """Dessine le schéma d'un écran de la solution.

    Args:
        ax: Axes matplotlib cible.
        x: Abscisse du bord gauche de l'écran.
        titre: Nom du module affiché dans la barre de titre.
        lignes: Éléments de l'écran, chacun décrit par un genre et un libellé.
            Genres acceptés : ``chip``, ``bulle``, ``reponse``, ``champ``,
            ``bouton``, ``note``.
    """
    largeur, hauteur = 3.0, 3.6
    ax.add_patch(
        FancyBboxPatch(
            (x, 0.35),
            largeur,
            hauteur,
            boxstyle="round,pad=0.03,rounding_size=0.10",
            linewidth=1.3,
            edgecolor=DARK,
            facecolor="white",
        )
    )
    # Barre de titre du module.
    ax.add_patch(
        FancyBboxPatch(
            (x, hauteur - 0.10),
            largeur,
            0.45,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            linewidth=0,
            facecolor=DARK,
        )
    )
    ax.text(x + 0.16, hauteur + 0.11, "OpenCacao", fontsize=7.6, color=ORANGE,
            fontweight="bold", va="center")
    ax.text(x + largeur - 0.16, hauteur + 0.11, titre, fontsize=8.4, color="white",
            fontweight="bold", va="center", ha="right")

    y = hauteur - 0.42
    for genre, libelle in lignes:
        if genre == "chip":
            ax.add_patch(
                FancyBboxPatch(
                    (x + 0.16, y - 0.22),
                    largeur - 0.32,
                    0.26,
                    boxstyle="round,pad=0.02,rounding_size=0.13",
                    linewidth=0.8,
                    edgecolor=ORANGE_LIGHT,
                    facecolor="white",
                )
            )
            ax.text(x + 0.26, y - 0.09, libelle, fontsize=6.8, color=GREY, va="center")
            y -= 0.34
        elif genre in {"bulle", "reponse"}:
            propre = libelle.split("\n")
            hauteur_bulle = 0.20 + 0.17 * len(propre)
            marge = 0.55 if genre == "bulle" else 0.16
            ax.add_patch(
                FancyBboxPatch(
                    (x + 0.16 + (marge if genre == "bulle" else 0), y - hauteur_bulle),
                    largeur - 0.32 - marge,
                    hauteur_bulle,
                    boxstyle="round,pad=0.02,rounding_size=0.10",
                    linewidth=0.9,
                    edgecolor=ORANGE if genre == "bulle" else ORANGE_LIGHT,
                    facecolor=ORANGE_PALE if genre == "bulle" else "white",
                )
            )
            ax.text(
                x + 0.26 + (marge if genre == "bulle" else 0),
                y - hauteur_bulle / 2,
                libelle,
                fontsize=6.8,
                color=DARK,
                va="center",
            )
            y -= hauteur_bulle + 0.12
        elif genre == "champ":
            ax.add_patch(
                FancyBboxPatch(
                    (x + 0.16, y - 0.30),
                    largeur - 0.32,
                    0.30,
                    boxstyle="round,pad=0.02,rounding_size=0.06",
                    linewidth=0.9,
                    edgecolor=GREY,
                    facecolor="white",
                )
            )
            ax.text(x + 0.26, y - 0.15, libelle, fontsize=6.8, color=GREY, va="center",
                    style="italic")
            y -= 0.40
        elif genre == "bouton":
            ax.add_patch(
                FancyBboxPatch(
                    (x + 0.16, y - 0.28),
                    largeur - 0.32,
                    0.28,
                    boxstyle="round,pad=0.02,rounding_size=0.08",
                    linewidth=0,
                    facecolor=ORANGE,
                )
            )
            ax.text(x + largeur / 2, y - 0.14, libelle, fontsize=7.2, color="white",
                    fontweight="bold", ha="center", va="center")
            y -= 0.38
        elif genre == "note":
            ax.text(x + 0.18, y - 0.12, libelle, fontsize=6.6, color=GREY, va="center",
                    style="italic")
            y -= 0.30


def figure_ecrans() -> Path:
    """Génère le schéma des trois écrans de la solution.

    Returns:
        Chemin du PNG produit.
    """
    fig = plt.figure(figsize=(9.4, 3.1))
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4.12)
    ax.axis("off")

    _ecran(ax, 0.25, "Le chat", [
        ("chip", "À quelle période tailler mes cacaoyers ?"),
        ("chip", "Comment réussir le séchage des fèves ?"),
        ("bulle", "Comment prévenir la pourriture\nbrune des cabosses ?"),
        ("reponse", "Réponse fondée sur les documents\nofficiels, sources citées.\n"
                    "Renvoi vers l'agent ANADER local."),
        ("champ", "Écrivez votre question sur le cacao…"),
    ])
    _ecran(ax, 3.50, "Ma parcelle", [
        ("champ", "Nom de la parcelle"),
        ("champ", "Localité : Daloa, Soubré, Divo…"),
        ("bouton", "Faire le tour de la parcelle (GPS)"),
        ("reponse", "Contour relevé, superficie calculée,\nphotos datées."),
        ("reponse", "Constat satellite : pas de déforestation\naprès le 31/12/2020."),
        ("note", "La photo est décrite, jamais diagnostiquée."),
    ])
    _ecran(ax, 6.75, "L'atelier", [
        ("champ", "Décrivez le document souhaité, en une phrase"),
        ("bouton", "Produire le document"),
        ("reponse", "Étude Word, tableau Excel\nou présentation PowerPoint."),
        ("reponse", "Manifeste joint : modèle, version,\ndocuments mobilisés, horodatage."),
        ("note", "Une section sans source est déclarée en lacune."),
    ])
    return _sauver(fig, "ecrans_solution.png")


def generer_visuels() -> dict[str, Path]:
    """Produit tous les PNG embarqués dans les deux documents.

    Returns:
        Dictionnaire clé logique vers chemin du PNG.
    """
    _style_matplotlib()
    return {
        "chiffres": figure_chiffres_cles(),
        "chaine": figure_chaine(),
        "architecture": figure_architecture(),
        "ecrans": figure_ecrans(),
        "performance": figure_performance(),
        "couverture": figure_couverture(),
        "repli": figure_repli(),
    }


# --------------------------------------------------------------------------- #
# Modèle de contenu commun aux deux moteurs de rendu                          #
# --------------------------------------------------------------------------- #


def contenu_dossier(visuels: dict[str, Path]) -> list[tuple[str, Any]]:
    """Décrit le dossier sous forme de blocs typés, indépendamment du format.

    Types de blocs : ``h1``, ``h2``, ``p``, ``pi`` (paragraphe secondaire),
    ``lead`` (intitulé en gras suivi d'un texte), ``puces``, ``tableau``,
    ``image``, ``encadre``, ``saut``.

    Args:
        visuels: Chemins des PNG produits par ``generer_visuels``.

    Returns:
        Liste ordonnée de blocs à rendre.
    """
    blocs: list[tuple[str, Any]] = []
    ajouter = blocs.append

    # ---------------------------------------------------------------- Résumé #
    ajouter(("h1", "Résumé exécutif"))
    ajouter((
        "p",
        "OpenCacao est un assistant de conseil agronomique dédié au cacao ivoirien. "
        "Il répond aux questions des producteurs et des techniciens en s'appuyant sur "
        "les documents officiels de la filière, il cite ses sources, il ne sort jamais "
        "du cacao et il renvoie systématiquement vers l'agent ANADER local pour tout ce "
        "qui relève de la décision de terrain. Il est en service, et ce dossier ne "
        "présente que des éléments mesurés.",
    ))
    ajouter(("image", (visuels["chiffres"], 5.2, "Le dossier en six chiffres, tous mesurés.")))
    ajouter(("h2", "Ce qu'il faut retenir"))
    ajouter((
        "puces",
        [
            ("Un socle de connaissance de la filière.", f"{VOLUME_SOURCE} de documents "
             f"officiels traités ont produit {EXTRAITS:,} extraits".replace(",", " ") +
             " vectorisés, interrogeables en quelques millisecondes."),
            ("Un modèle affiné pour le cacao.", "Ministral 3 8B Instruct à poids ouverts, "
             "affiné par LoRA 4 bits sur 10 000 paires questions/réponses cacao."),
            ("La souveraineté, sans exception.", "Aucun service d'IA externe dans le chemin "
             "de production. Ni OpenAI, ni Anthropic, ni Google. Les données des producteurs "
             "ne quittent pas l'infrastructure."),
            ("Sept agents spécialisés.", "Conseil documentaire, météo, prix officiel, "
             "réglementation, normes et certifications, satellite pour les alertes de "
             "déforestation, rédaction de livrables."),
            ("Une réponse en 2,3 secondes.", "Contre 38 secondes auparavant, et 123 tokens "
             "par seconde sur carte graphique contre 15 sur processeur. Mesures du 19/08/2026."),
            ("Une qualité logicielle vérifiable.", "1 655 tests automatisés et 99,57 % de "
             "couverture, branches comprises."),
            ("Une continuité de service éprouvée.", "Repli automatique constaté en production : "
             "50 secondes de détection, 75 secondes d'indisponibilité, aucune intervention."),
            ("Des garde-fous non négociables.", "L'outil refuse tout dosage phytosanitaire "
             "et oriente vers l'ANADER, traite le cacao et uniquement le cacao, refuse toute "
             "question médicale ou vétérinaire, et ne nomme jamais une maladie à partir "
             "d'une photo."),
            ("Des livrables traçables.", "Études et dossiers en Word, Excel et PowerPoint, "
             "chacun accompagné d'un manifeste de génération qui rend le document rejouable."),
        ],
    ))
    ajouter((
        "encadre",
        "Ce dossier n'avance ni taux d'adoption, ni nombre d'utilisateurs, ni retour sur "
        "investissement : ces éléments restent à établir. Tout ce qui est écrit ici est "
        "constaté.",
    ))
    ajouter(("saut", None))

    # ---------------------------------------------------------- Démonstrations #
    ajouter(("h1", "Neuf démonstrations exécutées en production"))
    ajouter((
        "p",
        "Les neuf scénarios ci-dessous ont tous été exécutés en production le 19 août 2026, "
        "sur le service en ligne. Ils sont reproductibles devant un tiers. Chacun tient en "
        "une phrase : c'est le matériau à utiliser face à un décideur.",
    ))
    ajouter(("h2", "Les refus qui protègent"))
    ajouter(("tableau", TABLEAU_SCENARIOS_REFUS))
    ajouter(("h2", "La lecture du terrain, sans fabrication"))
    ajouter(("tableau", TABLEAU_SCENARIOS_TERRAIN))
    ajouter(("h2", "La tenue du service et la traçabilité"))
    ajouter(("tableau", TABLEAU_SCENARIOS_SERVICE))
    ajouter((
        "encadre",
        "Ces neuf scénarios ont une caractéristique commune : dans chacun, la valeur de "
        "l'outil tient autant à ce qu'il refuse de dire qu'à ce qu'il dit.",
    ))
    ajouter(("saut", None))

    # ----------------------------------------------------------------- Besoin #
    ajouter(("h1", "1. Le besoin auquel OpenCacao répond"))
    ajouter((
        "p",
        "Le conseil agronomique du cacao ivoirien existe et il est de qualité. Il est "
        "produit, écrit et publié par le Conseil du Café-Cacao, l'ANADER, le CNRA, le "
        "FIRCA et leurs partenaires internationaux. Sa limite n'est pas son contenu : "
        "c'est sa disponibilité au moment précis où la question se pose, sur le support "
        "dont le producteur dispose réellement.",
    ))
    ajouter((
        "p",
        "OpenCacao ne remplace ni l'agent de terrain, ni l'institution qui produit la "
        "doctrine. Il rend le corpus officiel interrogeable en langage courant, en "
        "quelques secondes, et il ramène toujours l'échange vers l'agent ANADER local. "
        "C'est un outil de diffusion du conseil existant, pas une autorité concurrente.",
    ))
    ajouter(("h2", "Ce que l'outil résout, pour qui, et ce qui est vérifiable"))
    ajouter(("tableau", TABLEAU_VALEUR))
    ajouter(("saut", None))

    # ----------------------------------------------------------------- Agents #
    ajouter(("h1", "2. Ce que fait l'outil : sept agents spécialisés"))
    ajouter((
        "p",
        "Une question n'est pas traitée par un bloc unique. Elle traverse d'abord les "
        "garde-fous métier, puis elle est confiée à l'agent compétent. Chaque agent a un "
        "périmètre étroit et une limite écrite : c'est ce qui rend ses réponses "
        "défendables devant un technicien.",
    ))
    ajouter(("image", (visuels["chaine"], 6.6,
                       "Le chemin d'une question : garde-fous, routage, réponse sourcée.")))
    ajouter(("tableau", TABLEAU_AGENTS))
    ajouter(("saut", None))

    # ---------------------------------------------------------- Architecture #
    ajouter(("h1", "3. L'architecture technique, du serveur au producteur"))
    ajouter((
        "p",
        "Le schéma se lit de bas en haut, du calcul jusqu'à l'écran du producteur : "
        "le serveur qui fait tourner le modèle, le modèle lui-même, la connaissance "
        "vectorisée, l'orchestration des agents, puis les écrans. Chaque couche est "
        "opérée par le projet. Aucune ne délègue à un service d'IA externe, et c'est "
        "la seule façon de garantir que les données des producteurs ne circulent pas.",
    ))
    ajouter(("image", (visuels["architecture"], 6.6,
                       "Les cinq couches de la solution et la chaîne documentaire qui "
                       "alimente la base de connaissances.")))
    ajouter(("tableau", TABLEAU_ARCHITECTURE))
    ajouter(("h2", "Le calcul : le serveur GPU en service, le processeur en filet"))
    ajouter((
        "p",
        "Le modèle est servi par un serveur GPU dédié : c'est ce qui donne les 123 tokens "
        "par seconde et la réponse en 2,3 secondes. Si ce serveur devient indisponible, le "
        "service repart sur un repli en processeur, avec le même modèle au format quantifié : "
        "la réponse ralentit, elle ne s'arrête pas. Ce mécanisme n'est pas théorique, il a "
        "été éprouvé en production (voir la section consacrée à la continuité de service). "
        "Une troisième variante, locale, permet de servir le même modèle quantifié avec "
        "Ollama pour une démonstration hors infrastructure.",
    ))
    ajouter(("h2", "La console d'administration : trouver les documents, puis les vectoriser"))
    ajouter((
        "p",
        "La base de connaissances ne se remplit pas à la main. Une console "
        "d'administration, à accès réservé, exécute la chaîne documentaire de bout en bout :",
    ))
    ajouter((
        "puces",
        [
            ("Recherche automatique.", "La console explore une liste blanche de domaines "
             "officiels de la filière et repère les documents publiés qui ne sont pas encore "
             "connus. Le web ouvert n'est jamais suivi, et chaque adresse candidate est "
             "contrôlée avant tout téléchargement."),
            ("Téléchargement et magasin de documents.", "Les pièces retenues sont rapatriées "
             "sur le volume du projet, aux formats PDF, texte, Markdown, CSV et HTML."),
            ("Extraction, découpage, vectorisation.", "Le texte est extrait, découpé en "
             "extraits, vectorisé par le service d'embeddings, puis ajouté à l'index. L'ajout "
             "est additif : cette opération ne peut jamais réduire la base existante."),
            ("Préparation du prochain affinage.", "La console assemble, valide et dédoublonne "
             "le corpus retenu, prêt pour un entraînement sur GPU."),
            ("Accès réservé.", "La console est protégée par un formulaire de connexion, une "
             "session signée et une limitation de débit. Elle n'est pas destinée au public."),
        ],
    ))
    ajouter((
        "encadre",
        "Une limite connue, et assumée comme telle : les documents scannés, dont le texte "
        "n'est pas extractible, n'entrent pas encore dans la base. La reconnaissance de "
        "caractères est prévue, hors ligne, et seul le texte produit rejoindrait le service.",
    ))
    ajouter(("saut", None))

    # -------------------------------------------------------------- Trois écrans #
    ajouter(("h1", "4. La solution en trois écrans"))
    ajouter((
        "p",
        "Une seule adresse, trois usages. Il n'y a rien à installer : le service s'ouvre "
        "dans le navigateur d'un téléphone comme d'un ordinateur. Chaque écran a une "
        "fonction précise, et une seule.",
    ))
    ajouter(("image", (visuels["ecrans"], 6.6,
                       "Schémas des trois écrans : le chat, Ma parcelle, L'atelier.")))
    ajouter(("tableau", TABLEAU_MODULES))
    ajouter((
        "p",
        "Un quatrième écran existe, hors du parcours du producteur : la console "
        "d'administration décrite à la section précédente. Elle est réservée à l'équipe et "
        "aux experts métier habilités, parce qu'elle décide de ce qui entre dans la base de "
        "connaissances.",
    ))
    ajouter(("saut", None))

    # ------------------------------------------------------------------ Socle #
    ajouter(("h1", "5. Le socle de connaissance de la filière"))
    ajouter((
        "lead",
        ("Des sources officielles, et elles seules.", f"{VOLUME_SOURCE} de documents de la "
         "filière ont été traités pour construire ce qui fonctionne aujourd'hui. Ils ont "
         f"produit une base de {EXTRAITS:,} extraits vectorisés".replace(",", " ") +
         ", interrogeable en quelques millisecondes. Les organismes à l'origine de ces "
         "documents sont le Conseil du Café-Cacao, l'ANADER, le CNRA, le FIRCA, la FAO "
         "et l'ICCO."),
    ))
    ajouter((
        "lead",
        ("Un modèle affiné, pas un modèle généraliste habillé.", "Le moteur est Ministral 3 "
         "8B Instruct, un modèle à poids ouverts, affiné par LoRA 4 bits sur un corpus "
         "cacao ivoirien de 10 000 paires questions/réponses. L'affinage porte sur la "
         "manière de répondre à un producteur de cacao ; le corpus documentaire, lui, "
         "reste consultable et citable à chaque réponse."),
    ))
    ajouter((
        "lead",
        ("Une base qui grandit par la filière.", "La base n'est pas figée. Tout document "
         "officiel supplémentaire transmis par un expert métier peut être intégré, ce qui "
         "élargit d'autant le champ des questions traitées. La section de contact, en fin "
         "de dossier, décrit la marche à suivre."),
    ))
    ajouter(("saut", None))

    # ------------------------------------------------------------ Souveraineté #
    ajouter(("h1", "6. Souveraineté : où vont les données"))
    ajouter((
        "p",
        "C'est le point sur lequel un partenaire institutionnel doit pouvoir être "
        "catégorique. Aucun service d'IA externe n'intervient dans le chemin de "
        "production : ni OpenAI, ni Anthropic, ni Google. La question d'un producteur, sa "
        "localité, les photos de sa parcelle et le contour de son exploitation sont "
        "traités sur l'infrastructure du projet et n'en sortent pas.",
    ))
    ajouter((
        "puces",
        [
            ("Le modèle est hébergé par le projet.", "Poids ouverts, servi localement. "
             "Aucune requête sortante vers un fournisseur d'IA tiers."),
            ("Le corpus est constitué de documents de la filière.", "Il est identifiable, "
             "citable et remplaçable par ses ayants droit."),
            ("La chaîne est reproductible.", "L'affinage, la fusion du modèle et le "
             "déploiement sont scriptés et rejouables, sans dépendance à un service "
             "propriétaire."),
        ],
    ))
    ajouter((
        "p",
        "Cette contrainte a un coût technique, assumé : elle impose de maîtriser "
        "l'hébergement, la performance et la continuité de service. Les deux sections "
        "suivantes montrent que ce coût est tenu.",
    ))
    ajouter(("saut", None))

    # ------------------------------------------------------------- Garde-fous #
    ajouter(("h1", "7. Les garde-fous métier"))
    ajouter((
        "p",
        "Un assistant agronomique qui répond à tout est un risque pour la filière. "
        "OpenCacao refuse donc explicitement certaines demandes, et ces refus sont "
        "vérifiés en production. Ils ne sont pas un réglage : ce sont des règles non "
        "négociables du produit.",
    ))
    ajouter(("tableau", TABLEAU_GARDE_FOUS))
    ajouter((
        "encadre",
        "Chaque réponse porte un renvoi vers l'agent ANADER local. L'outil élargit "
        "l'accès au conseil ; il ne déplace jamais la responsabilité de la décision "
        "technique hors du réseau d'encadrement de la filière.",
    ))
    ajouter(("saut", None))

    # -------------------------------------------------- Vision et parcelle #
    ajouter(("h1", "8. Vision et parcelle : décrire, jamais diagnostiquer"))
    ajouter((
        "lead",
        ("La photo est décrite, elle n'est pas interprétée.", "Sur une photo de plantation, "
         "l'outil décrit objectivement ce qui est visible, et recoupe cette description "
         "avec la météo et l'historique de la parcelle. Il ne pose jamais de diagnostic. "
         "Le diagnostic reste à l'agent ANADER. C'est un choix assumé : un modèle "
         "généraliste est un bon descripteur et un mauvais diagnosticien, et un mauvais "
         "diagnostic coûte cher au producteur."),
    ))
    ajouter((
        "lead",
        ("La parcelle est documentée, pas déclarée.", "Le contour de la parcelle est relevé "
         "au GPS depuis un téléphone, la superficie est calculée, les photos sont datées, "
         "et un constat satellite atteste l'absence de déforestation postérieure au "
         "31 décembre 2020, date de référence du règlement européen sur la déforestation "
         "(EUDR)."),
    ))
    ajouter(("tableau", TABLEAU_VISION))
    ajouter((
        "p",
        "La limite est écrite dans le produit : l'outil fournit un constat satellite "
        "documenté et horodaté, il ne délivre aucune attestation de conformité EUDR. "
        "La certification reste du ressort des organismes habilités.",
    ))
    ajouter(("saut", None))

    # -------------------------------------------------------------- Livrables #
    ajouter(("h1", "9. Des livrables traçables et rejouables"))
    ajouter((
        "p",
        "Un partenaire ne juge pas un assistant sur une réponse à l'écran, mais sur les "
        "documents qu'il permet de produire. OpenCacao compose des études et des dossiers "
        "en Word, Excel et PowerPoint. Chaque document est accompagné d'un manifeste de "
        "génération, qui en fait une pièce vérifiable plutôt qu'un texte d'origine "
        "incertaine.",
    ))
    ajouter(("tableau", TABLEAU_MANIFESTE))
    ajouter((
        "encadre",
        "Règle de rédaction appliquée sans exception : une section pour laquelle aucune "
        "source n'est disponible est déclarée en lacune. Elle n'est jamais estimée.",
    ))
    ajouter(("saut", None))

    # ------------------------------------------------------------ Performance #
    ajouter(("h1", "10. Performance mesurée"))
    ajouter((
        "p",
        "Les valeurs ci-dessous ont été relevées le 19 août 2026, sur le même modèle et "
        "le même corpus, en changeant uniquement le matériel de service. Elles changent "
        "la nature de l'usage : à 38 secondes, on consulte un outil ; à 2,3 secondes, "
        "on dialogue avec lui.",
    ))
    ajouter(("image", (visuels["performance"], 6.6,
                       "Latence de bout en bout et débit de génération, mesures du 19/08/2026.")))
    ajouter(("tableau", TABLEAU_PERFORMANCE))
    ajouter(("saut", None))

    # ------------------------------------------------------------ Continuité #
    ajouter(("h1", "11. Continuité de service éprouvée"))
    ajouter((
        "p",
        "La disponibilité n'a pas été estimée sur plan : un repli automatique vers le "
        "matériel de secours a réellement eu lieu en production. La supervision a détecté "
        "la défaillance en 50 secondes, le service a été indisponible 75 secondes au "
        "total, et aucune intervention humaine n'a été nécessaire.",
    ))
    ajouter(("image", (visuels["repli"], 6.6,
                       "Décomposition des 75 secondes d'indisponibilité constatées en production.")))
    ajouter(("tableau", TABLEAU_REPLI))
    ajouter(("saut", None))

    # --------------------------------------------------------------- Qualité #
    ajouter(("h1", "12. Qualité logicielle vérifiable"))
    ajouter((
        "p",
        "Un outil qui porte du conseil agronomique doit pouvoir être audité. La base de "
        "code est couverte par 1 655 tests automatisés, pour une couverture de 99,57 %, "
        "branches comprises. Chaque règle de refus possède son propre test : un "
        "garde-fou n'est pas une intention, c'est un comportement vérifié à chaque "
        "modification du code.",
    ))
    ajouter(("image", (visuels["couverture"], 4.3,
                       "Couverture de tests, branches comprises.")))
    ajouter(("saut", None))

    # ------------------------------------------------------------- À établir #
    ajouter(("h1", "13. Ce qui reste à établir"))
    ajouter((
        "p",
        "La crédibilité d'un dossier tient autant à ce qu'il n'affirme pas. Les éléments "
        "suivants ne sont pas mesurés à ce jour et ne figurent volontairement pas dans ce "
        "document :",
    ))
    ajouter((
        "puces",
        [
            "le nombre d'utilisateurs et le taux d'adoption sur le terrain ;",
            "le gain de rendement ou de revenu chez les producteurs ;",
            "le retour sur investissement pour une coopérative ou un exportateur ;",
            "la part du corpus couvrant chaque organisme de la filière.",
        ],
    ))
    ajouter((
        "p",
        "Ces indicateurs supposent un déploiement suivi et un protocole de mesure "
        "convenu avec les institutions de la filière. Ils feront l'objet d'un relevé "
        "dédié, distinct de ce dossier.",
    ))

    # ------------------------------------------------------------ Conclusion #
    ajouter(("h1", "Conclusion"))
    ajouter((
        "p",
        "OpenCacao démontre qu'un conseil agronomique fondé sur les documents officiels "
        "de la filière peut être rendu immédiatement accessible, en français courant, "
        "sans confier une seule donnée de producteur à un fournisseur étranger. La "
        "démonstration n'est pas théorique : le service est en ligne, il répond en 2,3 "
        "secondes, il résiste à la panne d'un serveur en 75 secondes, et son "
        "comportement est verrouillé par 1 655 tests.",
    ))
    ajouter((
        "p",
        "L'outil a été conçu autour d'un principe simple : il élargit l'accès au conseil "
        "et il ramène toujours le producteur vers l'agent ANADER. Il ne prescrit pas, il "
        "ne diagnostique pas, il ne certifie pas. Ce périmètre étroit est ce qui le rend "
        "utilisable par la filière plutôt que concurrent de la filière.",
    ))
    ajouter((
        "p",
        "Ce qui manque aujourd'hui n'est pas de la technologie : c'est du corpus et du "
        "terrain. Chaque document officiel supplémentaire élargit immédiatement le champ "
        "des questions que l'outil sait traiter. C'est l'objet de la section suivante.",
    ))
    ajouter(("saut", None))

    # --------------------------------------------------------------- Contact #
    ajouter(("h1", "Contact et appel à contribution"))
    ajouter((
        "p",
        "La base de connaissances d'OpenCacao grandit par la filière, et seulement par "
        "elle. Si vous êtes agronome, technicien, chercheur, agent d'encadrement ou "
        "responsable qualité, vos documents ont une valeur directe : un guide technique, "
        "une fiche de vulgarisation, un référentiel de certification ou une note "
        "réglementaire transmis aujourd'hui devient une réponse sourcée demain.",
    ))
    ajouter(("h2", "Ce qui est utile à transmettre"))
    ajouter((
        "puces",
        [
            "les guides et fiches techniques cacao de votre organisme ;",
            "les notes réglementaires et les textes applicables à la filière ;",
            "les référentiels de normes et de certification, et leurs mises à jour ;",
            "les documents de vulgarisation destinés aux producteurs ;",
            "toute correction sur une réponse de l'outil que vous jugez inexacte ou "
            "incomplète.",
        ],
    ))
    ajouter(("h2", "Comment procéder"))
    ajouter((
        "puces",
        [
            ("1. Vous nous écrivez.", "Un courriel suffit, avec les documents en pièce "
             "jointe ou un lien de téléchargement, et le nom de l'organisme détenteur."),
            ("2. Nous vérifions les droits d'usage.", "Aucun document n'est intégré sans "
             "accord explicite de l'organisme qui en est à l'origine."),
            ("3. Le document est intégré et attribué.", "Il rejoint la base vectorisée et "
             "l'organisme est cité comme source dans chaque réponse qui s'appuie sur lui."),
            ("4. Vous constatez le résultat.", "Vous pouvez interroger le service en ligne "
             "et vérifier que vos documents sont bien mobilisés et correctement attribués."),
        ],
    ))
    ajouter(("h2", "Nous joindre"))
    ajouter((
        "tableau",
        (
            ["Objet", "Coordonnée"],
            [
                ["Contribution documentaire et partenariats",
                 "waopron@openlabconsulting.com"],
                ["Service en ligne, pour essayer l'outil",
                 "https://opencacao.openlabconsulting.com"],
                ["Éditeur", "OpenLab Consulting, Côte d'Ivoire"],
                ["Téléphone et adresse postale",
                 "à compléter par la direction de la communication avant diffusion"],
            ],
            [0.42, 0.58],
        ),
    ))
    ajouter((
        "pi",
        "OpenCacao est la démonstration technique du livre blanc « IA souveraine pour la "
        "Côte d'Ivoire » d'OpenLab Consulting.",
    ))
    return blocs


# --------------------------------------------------------------------------- #
# Rendu Word                                                                  #
# --------------------------------------------------------------------------- #


def _police(run: Any, *, size: float = 10.5, bold: bool = False,
            color: RGBColor = DOC_DARK, italic: bool = False) -> None:
    """Applique une police à un run Word.

    Args:
        run: Run python-docx à styler.
        size: Corps en points.
        bold: Graisse.
        color: Couleur d'encre.
        italic: Italique.
    """
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _fond(element: Any, hexcolor: str) -> None:
    """Applique un fond de couleur à un paragraphe ou une cellule Word.

    Args:
        element: Paragraphe (``p``) ou cellule (``tc``) python-docx.
        hexcolor: Couleur hexadécimale sans dièse.
    """
    proprietes = element._p.get_or_add_pPr() if hasattr(element, "_p") else element._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    proprietes.append(shd)


def _filet(paragraphe: Any, hexcolor: str = ORANGE_HEX, taille: str = "12") -> None:
    """Ajoute un filet horizontal sous un paragraphe Word.

    Args:
        paragraphe: Paragraphe python-docx.
        hexcolor: Couleur du filet.
        taille: Épaisseur en huitièmes de point.
    """
    proprietes = paragraphe._p.get_or_add_pPr()
    bordures = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), taille)
    bas.set(qn("w:space"), "4")
    bas.set(qn("w:color"), hexcolor)
    bordures.append(bas)
    proprietes.append(bordures)


def _pied_word(doc: Any) -> None:
    """Installe le pied de page avec le numéro de page dynamique.

    Args:
        doc: Document python-docx.
    """
    paragraphe = doc.sections[0].footer.paragraphs[0]
    paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _police(paragraphe.add_run(PIED + " · page "), size=8, color=DOC_GREY)
    run = paragraphe.add_run()
    debut, instruction, fin = (
        OxmlElement("w:fldChar"),
        OxmlElement("w:instrText"),
        OxmlElement("w:fldChar"),
    )
    debut.set(qn("w:fldCharType"), "begin")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    fin.set(qn("w:fldCharType"), "end")
    run._r.append(debut)
    run._r.append(instruction)
    run._r.append(fin)
    _police(run, size=8, color=DOC_GREY)


def _garde_word(doc: Any) -> None:
    """Compose la page de garde du document Word.

    Args:
        doc: Document python-docx.
    """
    if LOGO.exists():
        paragraphe = doc.add_paragraph()
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraphe.add_run().add_picture(str(LOGO), width=Inches(2.3))
    espace = doc.add_paragraph()
    espace.paragraph_format.space_after = Pt(60)

    paragraphe = doc.add_paragraph()
    paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphe.paragraph_format.space_after = Pt(4)
    _police(paragraphe.add_run(TITRE), size=46, bold=True, color=DOC_ORANGE)

    paragraphe = doc.add_paragraph()
    paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphe.paragraph_format.space_after = Pt(18)
    _filet(paragraphe)
    _police(paragraphe.add_run(SOUS_TITRE), size=14, color=DOC_DARK)

    paragraphe = doc.add_paragraph()
    paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphe.paragraph_format.space_after = Pt(60)
    _police(paragraphe.add_run(ACCROCHE), size=11, italic=True, color=DOC_GREY)

    bandeau = doc.add_paragraph()
    bandeau.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bandeau.paragraph_format.space_before = Pt(10)
    bandeau.paragraph_format.space_after = Pt(10)
    _fond(bandeau, ORANGE_HEX)
    _police(
        bandeau.add_run(
            "Souveraineté · sources officielles · garde-fous vérifiés · livrables traçables"
        ),
        size=11,
        bold=True,
        color=DOC_BLANC,
    )

    espace = doc.add_paragraph()
    espace.paragraph_format.space_after = Pt(80)

    paragraphe = doc.add_paragraph()
    paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _police(paragraphe.add_run(DATE_DOC + "\nOpenLab Consulting"), size=12, color=DOC_DARK)
    doc.add_page_break()


def _sommaire_word(doc: Any) -> None:
    """Compose le sommaire du document Word.

    Args:
        doc: Document python-docx.
    """
    _titre_word(doc, "Sommaire")
    tableau = doc.add_table(rows=0, cols=2)
    tableau.style = "Table Grid"
    tableau.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (titre, descriptif) in enumerate(SOMMAIRE):
        cellules = tableau.add_row().cells
        cellules[0].width = Inches(2.9)
        cellules[1].width = Inches(3.7)
        run = cellules[0].paragraphs[0].add_run(titre)
        _police(run, size=10, bold=True, color=DOC_ORANGE)
        run = cellules[1].paragraphs[0].add_run(descriptif)
        _police(run, size=10, color=DOC_GREY)
        if index % 2 == 0:
            for cellule in cellules:
                _fond(cellule, ORANGE_PALE_HEX)
    doc.add_page_break()


def _titre_word(doc: Any, texte: str) -> None:
    """Écrit un titre de niveau 1 souligné d'un filet orange.

    Args:
        doc: Document python-docx.
        texte: Libellé du titre.
    """
    paragraphe = doc.add_paragraph()
    paragraphe.paragraph_format.space_before = Pt(14)
    paragraphe.paragraph_format.space_after = Pt(10)
    paragraphe.paragraph_format.keep_with_next = True
    _filet(paragraphe)
    _police(paragraphe.add_run(texte), size=17, bold=True, color=DOC_ORANGE)


def _tableau_word(doc: Any, entetes: list[str], lignes: list[list[str]],
                  largeurs: list[float]) -> None:
    """Rend un tableau Word aux couleurs de la marque.

    Args:
        doc: Document python-docx.
        entetes: Libellés de la ligne d'en-tête.
        lignes: Contenu, une liste par ligne.
        largeurs: Largeurs relatives des colonnes (somme égale à 1).
    """
    utile = 6.6
    tableau = doc.add_table(rows=1, cols=len(entetes))
    tableau.style = "Table Grid"
    tableau.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, entete in enumerate(entetes):
        cellule = tableau.rows[0].cells[index]
        cellule.width = Inches(utile * largeurs[index])
        _fond(cellule, DARK_HEX)
        _police(cellule.paragraphs[0].add_run(entete), size=9.5, bold=True, color=DOC_BLANC)
    for rang, ligne in enumerate(lignes):
        cellules = tableau.add_row().cells
        for index, valeur in enumerate(ligne):
            cellule = cellules[index]
            cellule.width = Inches(utile * largeurs[index])
            if rang % 2 == 0:
                _fond(cellule, ORANGE_PALE_HEX)
            _police(cellule.paragraphs[0].add_run(valeur), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def rendre_word(blocs: list[tuple[str, Any]]) -> Path:
    """Rend le dossier au format Word.

    Args:
        blocs: Modèle de contenu produit par ``contenu_dossier``.

    Returns:
        Chemin du .docx écrit.
    """
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)

    _pied_word(doc)
    _garde_word(doc)
    _sommaire_word(doc)

    for genre, valeur in blocs:
        if genre == "h1":
            _titre_word(doc, valeur)
        elif genre == "h2":
            paragraphe = doc.add_paragraph()
            paragraphe.paragraph_format.space_before = Pt(10)
            paragraphe.paragraph_format.space_after = Pt(5)
            paragraphe.paragraph_format.keep_with_next = True
            _police(paragraphe.add_run(valeur), size=12.5, bold=True, color=DOC_DARK)
        elif genre in {"p", "pi"}:
            paragraphe = doc.add_paragraph()
            paragraphe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraphe.paragraph_format.space_after = Pt(7)
            _police(
                paragraphe.add_run(valeur),
                size=10.5 if genre == "p" else 9.5,
                italic=genre == "pi",
                color=DOC_DARK if genre == "p" else DOC_GREY,
            )
        elif genre == "lead":
            intitule, texte = valeur
            paragraphe = doc.add_paragraph()
            paragraphe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraphe.paragraph_format.space_after = Pt(7)
            _police(paragraphe.add_run(intitule + " "), size=10.5, bold=True, color=DOC_ORANGE)
            _police(paragraphe.add_run(texte), size=10.5)
        elif genre == "puces":
            for item in valeur:
                paragraphe = doc.add_paragraph(style="List Bullet")
                paragraphe.paragraph_format.space_after = Pt(3)
                if isinstance(item, tuple):
                    _police(paragraphe.add_run(item[0] + " "), size=10.5, bold=True,
                            color=DOC_ORANGE)
                    _police(paragraphe.add_run(item[1]), size=10.5)
                else:
                    _police(paragraphe.add_run(item), size=10.5)
        elif genre == "tableau":
            _tableau_word(doc, *valeur)
        elif genre == "image":
            chemin, largeur, legende = valeur
            paragraphe = doc.add_paragraph()
            paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraphe.paragraph_format.space_before = Pt(6)
            paragraphe.paragraph_format.space_after = Pt(3)
            paragraphe.add_run().add_picture(str(chemin), width=Inches(largeur))
            paragraphe = doc.add_paragraph()
            paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraphe.paragraph_format.space_after = Pt(10)
            _police(paragraphe.add_run(legende), size=8.5, italic=True, color=DOC_GREY)
        elif genre == "encadre":
            paragraphe = doc.add_paragraph()
            paragraphe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraphe.paragraph_format.space_before = Pt(8)
            paragraphe.paragraph_format.space_after = Pt(10)
            paragraphe.paragraph_format.left_indent = Pt(10)
            paragraphe.paragraph_format.right_indent = Pt(10)
            _fond(paragraphe, ORANGE_PALE_HEX)
            _police(paragraphe.add_run(valeur), size=10.5, italic=True, color=DOC_DARK)
        elif genre == "saut":
            doc.add_page_break()

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


# --------------------------------------------------------------------------- #
# Rendu PDF                                                                   #
# --------------------------------------------------------------------------- #

LARGEUR_UTILE = A4[0] - 2 * 2.4 * cm


def _styles_pdf() -> dict[str, ParagraphStyle]:
    """Construit la feuille de styles platypus aux couleurs de la marque.

    Returns:
        Dictionnaire de styles indexés par nom logique.
    """
    base = getSampleStyleSheet()["Normal"]
    commun = {"fontName": "Helvetica", "textColor": PDF_DARK}
    return {
        "corps": ParagraphStyle("corps", parent=base, fontSize=10.5, leading=15.2,
                                alignment=TA_JUSTIFY, spaceAfter=7, **commun),
        "corps_sec": ParagraphStyle("corps_sec", parent=base, fontSize=9.5, leading=13.5,
                                    alignment=TA_JUSTIFY, spaceAfter=7,
                                    fontName="Helvetica-Oblique", textColor=PDF_GREY),
        "h1": ParagraphStyle("h1", parent=base, fontSize=17, leading=21, spaceBefore=6,
                             spaceAfter=4, fontName="Helvetica-Bold", textColor=PDF_ORANGE),
        "h2": ParagraphStyle("h2", parent=base, fontSize=12.5, leading=16, spaceBefore=10,
                             spaceAfter=5, fontName="Helvetica-Bold", textColor=PDF_DARK),
        "puce": ParagraphStyle("puce", parent=base, fontSize=10.5, leading=14.6,
                               alignment=TA_JUSTIFY, spaceAfter=4, leftIndent=14,
                               bulletIndent=3, **commun),
        "legende": ParagraphStyle("legende", parent=base, fontSize=8.5, leading=11,
                                  alignment=TA_CENTER, spaceAfter=10,
                                  fontName="Helvetica-Oblique", textColor=PDF_GREY),
        "encadre": ParagraphStyle("encadre", parent=base, fontSize=10.5, leading=15,
                                  alignment=TA_JUSTIFY, fontName="Helvetica-Oblique",
                                  textColor=PDF_DARK, leftIndent=8, rightIndent=8,
                                  spaceBefore=6, spaceAfter=6),
        "cellule": ParagraphStyle("cellule", parent=base, fontSize=8.8, leading=11.6, **commun),
        "cellule_tete": ParagraphStyle("cellule_tete", parent=base, fontSize=8.8, leading=11.6,
                                       fontName="Helvetica-Bold", textColor=colors.white),
        "garde_titre": ParagraphStyle("garde_titre", parent=base, fontSize=46, leading=52,
                                      alignment=TA_CENTER, fontName="Helvetica-Bold",
                                      textColor=PDF_ORANGE),
        "garde_sous": ParagraphStyle("garde_sous", parent=base, fontSize=14, leading=19,
                                     alignment=TA_CENTER, fontName="Helvetica",
                                     textColor=PDF_DARK, spaceAfter=10),
        "garde_accroche": ParagraphStyle("garde_accroche", parent=base, fontSize=11, leading=15,
                                         alignment=TA_CENTER, fontName="Helvetica-Oblique",
                                         textColor=PDF_GREY),
        "garde_date": ParagraphStyle("garde_date", parent=base, fontSize=12, leading=17,
                                     alignment=TA_CENTER, fontName="Helvetica",
                                     textColor=PDF_DARK),
        "bandeau": ParagraphStyle("bandeau", parent=base, fontSize=11, leading=15,
                                  alignment=TA_CENTER, fontName="Helvetica-Bold",
                                  textColor=colors.white),
        "somm_titre": ParagraphStyle("somm_titre", parent=base, fontSize=10, leading=13,
                                     fontName="Helvetica-Bold", textColor=PDF_ORANGE),
        "somm_desc": ParagraphStyle("somm_desc", parent=base, fontSize=10, leading=13,
                                    fontName="Helvetica", textColor=PDF_GREY),
    }


def _pied_pdf(canevas: Any, doc: Any) -> None:
    """Dessine le pied de page du PDF, sauf sur la page de garde.

    Args:
        canevas: Canevas reportlab de la page courante.
        doc: Document platypus en cours de rendu.
    """
    if canevas.getPageNumber() == 1:
        return
    canevas.saveState()
    canevas.setStrokeColor(PDF_LIGNE)
    canevas.setLineWidth(0.6)
    canevas.line(2.4 * cm, 1.55 * cm, A4[0] - 2.4 * cm, 1.55 * cm)
    canevas.setFont("Helvetica", 7.5)
    canevas.setFillColor(PDF_GREY)
    canevas.drawString(2.4 * cm, 1.15 * cm, PIED)
    canevas.drawRightString(A4[0] - 2.4 * cm, 1.15 * cm, f"page {canevas.getPageNumber()}")
    canevas.restoreState()


def _image_pdf(chemin: Path, largeur_pouces: float) -> Image:
    """Crée une image platypus mise à l'échelle en conservant ses proportions.

    Args:
        chemin: Chemin du PNG.
        largeur_pouces: Largeur cible en pouces, identique à celle du Word.

    Returns:
        Flowable image prêt à être ajouté au récit.
    """
    from PIL import Image as PILImage

    with PILImage.open(chemin) as source:
        ratio = source.height / source.width
    largeur = largeur_pouces * inch
    return Image(str(chemin), width=largeur, height=largeur * ratio)


def _tableau_pdf(entetes: list[str], lignes: list[list[str]], largeurs: list[float],
                 styles: dict[str, ParagraphStyle]) -> Table:
    """Rend un tableau PDF aux couleurs de la marque.

    Args:
        entetes: Libellés de la ligne d'en-tête.
        lignes: Contenu, une liste par ligne.
        largeurs: Largeurs relatives des colonnes (somme égale à 1).
        styles: Feuille de styles issue de ``_styles_pdf``.

    Returns:
        Flowable tableau prêt à être ajouté au récit.
    """
    donnees = [[Paragraph(entete, styles["cellule_tete"]) for entete in entetes]]
    donnees += [[Paragraph(valeur, styles["cellule"]) for valeur in ligne] for ligne in lignes]
    tableau = Table(
        donnees,
        colWidths=[LARGEUR_UTILE * part for part in largeurs],
        repeatRows=1,
        hAlign="CENTER",
    )
    commandes = [
        ("BACKGROUND", (0, 0), (-1, 0), PDF_DARK),
        ("GRID", (0, 0), (-1, -1), 0.5, PDF_LIGNE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]
    # Zébrage : une ligne sur deux en orange très pâle.
    for rang in range(1, len(donnees)):
        if rang % 2 == 1:
            commandes.append(("BACKGROUND", (0, rang), (-1, rang), PDF_ORANGE_PALE))
    tableau.setStyle(TableStyle(commandes))
    return tableau


def _garde_pdf(styles: dict[str, ParagraphStyle]) -> list[Any]:
    """Compose la page de garde du PDF.

    Args:
        styles: Feuille de styles issue de ``_styles_pdf``.

    Returns:
        Liste de flowables constituant la page de garde.
    """
    recit: list[Any] = [Spacer(1, 1.1 * cm)]
    if LOGO.exists():
        logo = _image_pdf(LOGO, 2.3)
        logo.hAlign = "CENTER"
        recit.append(logo)
    recit += [
        Spacer(1, 2.6 * cm),
        Paragraph(TITRE, styles["garde_titre"]),
        Spacer(1, 0.25 * cm),
        HRFlowable(width="70%", thickness=1.4, color=PDF_ORANGE, hAlign="CENTER",
                   spaceAfter=10),
        Paragraph(SOUS_TITRE, styles["garde_sous"]),
        Paragraph(ACCROCHE, styles["garde_accroche"]),
        Spacer(1, 2.2 * cm),
    ]
    bandeau = Table(
        [[Paragraph(
            "Souveraineté &middot; sources officielles &middot; garde-fous vérifiés "
            "&middot; livrables traçables",
            styles["bandeau"],
        )]],
        colWidths=[LARGEUR_UTILE],
    )
    bandeau.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), PDF_ORANGE),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    recit += [
        bandeau,
        Spacer(1, 3.4 * cm),
        Paragraph(DATE_DOC + "<br/>OpenLab Consulting", styles["garde_date"]),
        PageBreak(),
    ]
    return recit


def _sommaire_pdf(styles: dict[str, ParagraphStyle]) -> list[Any]:
    """Compose le sommaire du PDF.

    Args:
        styles: Feuille de styles issue de ``_styles_pdf``.

    Returns:
        Liste de flowables constituant le sommaire.
    """
    donnees = [
        [Paragraph(titre, styles["somm_titre"]), Paragraph(descriptif, styles["somm_desc"])]
        for titre, descriptif in SOMMAIRE
    ]
    tableau = Table(donnees, colWidths=[LARGEUR_UTILE * 0.44, LARGEUR_UTILE * 0.56],
                    hAlign="CENTER")
    commandes = [
        ("GRID", (0, 0), (-1, -1), 0.5, PDF_LIGNE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]
    for rang in range(len(donnees)):
        if rang % 2 == 0:
            commandes.append(("BACKGROUND", (0, rang), (-1, rang), PDF_ORANGE_PALE))
    tableau.setStyle(TableStyle(commandes))
    return [
        Paragraph("Sommaire", styles["h1"]),
        HRFlowable(width="100%", thickness=1.2, color=PDF_ORANGE, spaceAfter=12),
        tableau,
        PageBreak(),
    ]


def rendre_pdf(blocs: list[tuple[str, Any]]) -> Path:
    """Rend le dossier au format PDF.

    Args:
        blocs: Modèle de contenu produit par ``contenu_dossier``.

    Returns:
        Chemin du .pdf écrit.
    """
    styles = _styles_pdf()
    OUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=2.4 * cm,
        rightMargin=2.4 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.1 * cm,
        title="OpenCacao, dossier de présentation",
        author="OpenLab Consulting",
        subject=SOUS_TITRE,
    )
    recit: list[Any] = _garde_pdf(styles) + _sommaire_pdf(styles)

    for genre, valeur in blocs:
        if genre == "h1":
            recit.append(
                KeepTogether(
                    [
                        Paragraph(valeur, styles["h1"]),
                        HRFlowable(width="100%", thickness=1.2, color=PDF_ORANGE, spaceAfter=10),
                    ]
                )
            )
        elif genre == "h2":
            recit.append(Paragraph(valeur, styles["h2"]))
        elif genre == "p":
            recit.append(Paragraph(valeur, styles["corps"]))
        elif genre == "pi":
            recit.append(Paragraph(valeur, styles["corps_sec"]))
        elif genre == "lead":
            intitule, texte = valeur
            recit.append(
                Paragraph(
                    f'<font color="#{ORANGE_HEX}"><b>{intitule}</b></font> {texte}',
                    styles["corps"],
                )
            )
        elif genre == "puces":
            for item in valeur:
                texte = (
                    f'<font color="#{ORANGE_HEX}"><b>{item[0]}</b></font> {item[1]}'
                    if isinstance(item, tuple)
                    else item
                )
                recit.append(Paragraph(texte, styles["puce"], bulletText="•"))
            recit.append(Spacer(1, 4))
        elif genre == "tableau":
            entetes, lignes, largeurs = valeur
            recit.append(Spacer(1, 3))
            recit.append(_tableau_pdf(entetes, lignes, largeurs, styles))
            recit.append(Spacer(1, 9))
        elif genre == "image":
            chemin, largeur, legende = valeur
            image = _image_pdf(chemin, largeur)
            image.hAlign = "CENTER"
            recit.append(
                KeepTogether([Spacer(1, 4), image, Spacer(1, 4),
                              Paragraph(legende, styles["legende"])])
            )
        elif genre == "encadre":
            encadre = Table(
                [[Paragraph(valeur, styles["encadre"])]],
                colWidths=[LARGEUR_UTILE],
                hAlign="CENTER",
            )
            encadre.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), PDF_ORANGE_PALE),
                        ("LINEBEFORE", (0, 0), (0, -1), 2.4, PDF_ORANGE),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ]
                )
            )
            recit += [Spacer(1, 4), encadre, Spacer(1, 10)]
        elif genre == "saut":
            recit.append(PageBreak())

    doc.build(recit, onFirstPage=_pied_pdf, onLaterPages=_pied_pdf)
    return OUT_PDF


# --------------------------------------------------------------------------- #
# Point d'entrée                                                              #
# --------------------------------------------------------------------------- #


def build() -> tuple[Path, Path]:
    """Génère les visuels puis les deux documents.

    Returns:
        Couple (chemin du .docx, chemin du .pdf).
    """
    visuels = generer_visuels()
    blocs = contenu_dossier(visuels)
    return rendre_word(blocs), rendre_pdf(blocs)


if __name__ == "__main__":
    chemin_docx, chemin_pdf = build()
    print(f"OK -> {chemin_docx}")
    print(f"OK -> {chemin_pdf}")
