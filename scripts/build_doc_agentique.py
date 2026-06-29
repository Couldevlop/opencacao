"""Construit la documentation Word du socle agentique V3 d'OpenCacao.

Document pédagogique ET d'architecture : page de garde, vue d'ensemble souveraine,
architecture en couches, un chapitre par brique (contrat, registre, routeur,
orchestrateur, squelette d'agent, tool use, réplication, synthèse multi-agents),
câblage derrière flag, recette d'extension « ajouter un agent en 4 étapes »,
méthodologie TDD, garde-fous et glossaire. Chaque chapitre suit la trame
« Concept → Décisions de conception → Modèle mental → Code & tests ».

Usage : python scripts/build_doc_agentique.py
Sortie : docs/Documentation_Socle_Agentique_V3.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "docs" / "OPENLAB.png"
CACAO = ROOT / "docs" / "Cacao.png"
OUT = ROOT / "docs" / "Documentation_Socle_Agentique_V3.docx"

OR = RGBColor(0xEA, 0x5B, 0x13)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GREY = RGBColor(0x60, 0x60, 0x60)
CODEBG = "F2F2F2"


def _set_font(run, size=11, bold=False, color=DARK, italic=False, name=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    if name:
        run.font.name = name


def heading(doc, text, level=1, color=OR):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    r.font.color.rgb = color
    return p


def para(doc, text, *, size=11, italic=False, color=DARK, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    _set_font(p.add_run(text), size=size, italic=italic, color=color)
    return p


def lead(doc, label, text):
    """Paragraphe « Intitulé. texte » avec l'intitulé en gras coloré."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _set_font(p.add_run(label + " "), bold=True, color=OR)
    _set_font(p.add_run(text))
    return p


def bullets(doc, items, color=DARK):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            _set_font(p.add_run(it[0] + " "), bold=True, color=OR)
            _set_font(p.add_run(it[1]), color=color)
        else:
            _set_font(p.add_run(it), color=color)


def numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        if isinstance(it, tuple):
            _set_font(p.add_run(it[0] + " "), bold=True, color=OR)
            _set_font(p.add_run(it[1]))
        else:
            _set_font(p.add_run(it))


def _shade(p, color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def code(doc, text):
    """Bloc de code monospace sur fond gris clair."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(8)
    _shade(p, CODEBG)
    for i, line in enumerate(text.strip("\n").split("\n")):
        run = p.add_run(("\n" if i else "") + line)
        _set_font(run, size=9, color=DARK, name="Consolas")


def quote(doc, text):
    """Encadré « modèle mental » en italique coloré."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    _set_font(p.add_run("➤ " + text), italic=True, color=OR)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        r = t.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val
            for pr in cells[j].paragraphs:
                for rn in pr.runs:
                    rn.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def hrule(p, color="EA5B13"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def footer(doc):
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("OpenLab Consulting — OpenCacao · Socle agentique V3 · "), size=8, color=GREY)
    run = p.add_run()
    fld1, instr, fld2 = OxmlElement("w:fldChar"), OxmlElement("w:instrText"), OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    _set_font(run, size=8, color=GREY)


def cover(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(2.2))
    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("Plateforme agentique V3"), size=30, bold=True, color=OR)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("Architecture & cours d'IA agentique"), size=16, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(
        p.add_run("Orchestrateur souverain · registre extensible · agents spécialisés · tool use"),
        size=11,
        italic=True,
        color=GREY,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    if CACAO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(CACAO), width=Inches(2.6))
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(
        p.add_run("OpenCacao-8B — IA souveraine pour la Côte d'Ivoire\nOpenLab Consulting"),
        size=11,
        color=DARK,
    )
    doc.add_page_break()


def chapter(doc, num, titre, concept, decisions, mental, code_blocks=None, tests=None, extra=None):
    """Chapitre standard : Concept → Décisions → Modèle mental → Code & tests."""
    heading(doc, f"{num}. {titre}", level=1)
    heading(doc, "Le concept", level=3)
    for c in concept:
        para(doc, c)
    heading(doc, "Les décisions de conception", level=3)
    bullets(doc, decisions)
    if extra:
        extra(doc)
    heading(doc, "Modèle mental", level=3)
    quote(doc, mental)
    if code_blocks:
        heading(doc, "Code", level=3)
        for label, snippet in code_blocks:
            if label:
                para(doc, label, size=10, italic=True, color=GREY, space_after=2)
            code(doc, snippet)
    if tests:
        heading(doc, "Ce que les tests verrouillent", level=3)
        bullets(doc, tests)
    doc.add_page_break()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    footer(doc)

    cover(doc)

    # --- Préambule ---
    heading(doc, "Préambule — comment lire ce document", level=1)
    para(
        doc,
        "Ce document a deux fonctions indissociables : c'est la documentation "
        "d'architecture du socle agentique V3 d'OpenCacao, et c'est un cours d'IA "
        "agentique. Chaque brique logicielle y est expliquée selon la même trame :",
    )
    bullets(
        doc,
        [
            ("Le concept —", "le pattern agentique en jeu, en langage clair."),
            ("Les décisions de conception —", "pourquoi on a tranché ainsi (l'expertise est là)."),
            ("Le modèle mental —", "la phrase à retenir."),
            ("Code & tests —", "le code réel et ce que les tests garantissent."),
        ],
    )
    para(
        doc,
        "L'ordre des chapitres EST la progression pédagogique : du contrat (ce qu'est "
        "un agent) jusqu'à la synthèse multi-agents (plusieurs agents qui coopèrent). "
        "Lis dans l'ordre.",
    )
    para(
        doc,
        "Méthode de construction : TDD strict (test rouge → code minimal → test vert → "
        "commit), couverture ≥ 97 %, inférence et réseau mockés en test. Aucun appel à "
        "un LLM tiers en production : souveraineté.",
        italic=True,
        color=GREY,
    )
    doc.add_page_break()

    # --- Vue d'ensemble ---
    heading(doc, "Vue d'ensemble", level=1)
    para(
        doc,
        "OpenCacao-8B est un assistant de conseil agronomique pour les producteurs de "
        "cacao ivoiriens, basé sur Ministral-3-8B affiné par LoRA, servi par une API "
        "FastAPI. La V2 répond en un appel (RAG + garde-fous + cache). La V3 introduit "
        "une plateforme agentique : un orchestrateur route chaque requête vers des agents "
        "spécialisés (RAG, Météo, Prix, Reporting) enregistrés dans un registre dynamique.",
    )
    lead(
        doc,
        "Objectif d'architecture.",
        "un framework EXTENSIBLE à 10+ agents. Ajouter l'agent n°5..n°11 = écrire une "
        "classe conforme au contrat + l'enregistrer. Aucune modification de "
        "l'orchestrateur, du registre ou du routeur.",
    )
    lead(
        doc,
        "Rétrocompatibilité.",
        "la plateforme est livrée derrière un flag (agents_enabled, OFF par défaut). Tant "
        "qu'il est OFF, le chemin V2 (ConseilService) reste seul en production.",
    )
    heading(doc, "Le flux d'une requête", level=3)
    code(
        doc,
        """
requête (question, langue, historique, ip)
        │
        ▼
┌──────────────────  ORCHESTRATEUR  ──────────────────┐
│ 1. fil_ancre        (anti-dérive multi-tours)        │
│ 2. garde-fou ENTRÉE (cacao-only) ──► refus ──────────┼─► Conseil (ANADER)
│ 3. ROUTEUR          (qui répond ? score peut_traiter)│
│ 4. rate-limit       (avant inférence, après routage) │
│ 5. dispatch ─► AGENT ─► OUTIL (météo/prix) ─► LLM     │
│ 6. garde-fou SORTIE (vérifie la génération)          │
│ 7. journalisation   (trace + interaction_id)         │
└──────────────────────────────────────────────────────┘
        │
        ▼
Conseil (réponse + sources + confiance + disclaimer ANADER)
""",
    )
    heading(doc, "Architecture en couches (clean architecture)", level=3)
    table(
        doc,
        ["Couche", "Fichiers", "Rôle"],
        [
            ["Domaine (pur)", "domain/agents.py", "Contrat : AgentRequete/Reponse/Port/Outil. Aucune dépendance framework."],
            ["Application (pur)", "application/registre,routage,orchestrateur,contexte.py", "Orchestration testable sans réseau."],
            ["Adaptateurs", "services/agents/*", "Agents concrets (RAG, Météo, Prix, Reporting)."],
            ["Adaptateurs", "services/outils/*", "Outils concrets (météo, prix) derrière un port mockable."],
        ],
    )
    para(
        doc,
        "La frontière « contrat/orchestration pure » ↔ « adaptateurs concrets » est ce qui "
        "rend la plateforme extensible : un nouvel agent n'est qu'un nouvel adaptateur.",
        italic=True,
        color=GREY,
    )
    doc.add_page_break()

    # --- Chapitres ---
    chapter(
        doc,
        1,
        "Le contrat d'agent — domain/agents.py",
        concept=[
            "Un agent est une capacité bornée derrière une interface stable : il déclare "
            "ce qu'il sait faire (routage), reçoit une requête normalisée, rend une réponse "
            "normalisée. Le reste du système ne connaît QUE cette interface, jamais "
            "l'implémentation. C'est l'inversion de dépendance (le « D » de SOLID) appliquée "
            "à l'agentique.",
        ],
        decisions=[
            ("dataclass(frozen=True) —", "requêtes/réponses immuables. Un agent ne peut pas modifier par surprise une donnée qu'un autre lira. Pilier de fiabilité en async."),
            ("Protocol (typage structurel) —", "un agent est conforme parce qu'il a les bonnes méthodes, pas par héritage. On ne piège pas les futurs agents dans une classe mère."),
            ("peut_traiter() -> float (0..1) —", "chaque agent s'auto-évalue. Il se décrit ; le routeur décide. Déterministe (mots-clés) : explicable, testable, souverain."),
            ("invoquer(**kwargs) —", "absorbe la variabilité des outils (météo prend localite, prix ne prend rien). Un bon contrat anticipe la diversité des implémentations."),
        ],
        mental="Le contrat est la constitution de la plateforme. Tout en dépend ; lui ne dépend de rien. Ajouter l'agent n°11 = classe conforme + enregistrement. Zéro refactor.",
        code_blocks=[
            (
                "Les quatre pièces du contrat (extrait) :",
                """
@dataclass(frozen=True)
class AgentRequete:
    question: str
    langue: Langue
    fil_ancre: str
    client_ip: str
    historique: list[dict[str, str]] = field(default_factory=list)

@runtime_checkable
class AgentPort(Protocol):
    nom: str
    description: str
    mots_cles: tuple[str, ...]
    async def peut_traiter(self, requete: AgentRequete) -> float: ...
    async def traiter(self, requete: AgentRequete) -> AgentReponse: ...
""",
            ),
        ],
        tests=[
            "Un agent factice SANS héritage est reconnu isinstance(..., AgentPort) → preuve du typage structurel.",
            "Modifier une AgentRequete lève FrozenInstanceError → preuve de l'immuabilité.",
            "peut_traiter renvoie bien un score différencié selon la question.",
        ],
    )

    chapter(
        doc,
        2,
        "Le registre dynamique — application/registre.py",
        concept=[
            "Un annuaire d'agents : on enregistre des instances, on les retrouve par nom ou "
            "énumération. C'est le point d'extension n°1 : il rend la plateforme « ouverte à "
            "l'extension, fermée à la modification » (le « O » de SOLID).",
        ],
        decisions=[
            ("Refus des doublons (ValueError) —", "deux agents nommés « meteo » = l'un écrase l'autre silencieusement = bug. On l'interdit."),
            ("Énumération stable (tous/noms) —", "le routeur balaie les agents enregistrés."),
            ("Observabilité —", "chaque enregistrement est journalisé (structlog), utile en prod pour savoir quels agents sont montés."),
        ],
        mental="Le registre est la prise électrique du framework. Brancher un agent suffit à le rendre routable ; rien d'autre ne bouge.",
        code_blocks=[
            (
                None,
                """
class RegistreAgents:
    def enregistrer(self, agent: AgentPort) -> None:
        if agent.nom in self._agents:
            raise ValueError(f"Agent « {agent.nom} » déjà enregistré")
        self._agents[agent.nom] = agent
        logger.info("agent_enregistre", agent=agent.nom)
""",
            ),
        ],
        tests=[
            "Enregistrer puis obtenir rend la même instance.",
            "Un nom dupliqué est rejeté.",
            "tous()/noms() exposent l'ensemble enregistré.",
        ],
    )

    chapter(
        doc,
        3,
        "Le routeur d'intention — application/routage.py",
        concept=[
            "« Qui doit répondre ? » Chaque agent s'auto-évalue (peut_traiter) ; le routeur "
            "classe par score décroissant et coupe sous un seuil. C'est la graine du planner "
            "des architectures multi-agents (ReAct, plan-and-execute), ici en version plate "
            "(un tour).",
        ],
        decisions=[
            ("Déterministe d'abord —", "aucun appel LLM pour router : explicable, testable, souverain. L'interface ne changera pas si on bascule plus tard vers un routage sémantique (embeddings)."),
            ("Un classement, pas un seul gagnant —", "certaines requêtes mobilisent plusieurs agents. Le routeur renvoie une liste ordonnée ; l'orchestrateur décide combien activer."),
        ],
        mental="Le routeur note, il ne décide pas seul. La décision finale (1 ou N agents, repli) appartient à l'orchestrateur.",
        code_blocks=[
            (
                None,
                """
async def classer(self, requete) -> list[tuple[AgentPort, float]]:
    scores = []
    for agent in self._registre.tous():
        score = await agent.peut_traiter(requete)
        if score >= self._seuil:
            scores.append((agent, score))
    scores.sort(key=lambda paire: paire[1], reverse=True)
    return scores
""",
            ),
        ],
        tests=[
            "Le classement est trié par score décroissant et exclut les scores sous le seuil.",
            "meilleur() renvoie le plus pertinent, ou None si tous sont sous le seuil.",
        ],
    )

    chapter(
        doc,
        4,
        "L'orchestrateur — application/orchestrateur.py (le cœur)",
        concept=[
            "C'est le control plane (plan de contrôle) / la boucle d'agent. Tout le reste "
            "exécute ; lui décide : qui agit, dans quel ordre, sous quelles contraintes. "
            "Équivalent agentique de ConseilService. L'ordre de ses 7 étapes encode la "
            "sécurité et l'équité.",
        ],
        decisions=[
            ("Garde-fous CENTRALISÉS, pas par agent —", "point d'application unique de la politique. Le filtre « cacao uniquement » ne peut pas être oublié sur un futur agent. Souveraineté structurelle."),
            ("Défense en profondeur (entrée ET sortie) —", "l'entrée bloque la demande interdite (sur le fil ancré → pas de contournement multi-tours) ; la sortie inspecte ce que l'agent a réellement généré. On ne fait jamais confiance à la sortie d'un LLM sans la vérifier."),
            ("Rate-limit après le routage, avant l'inférence —", "un refus ne coûte pas de génération CPU (~38 s) → il ne doit pas consommer le quota. On ne facture que le travail coûteux. Équité."),
            ("Repli systématique —", "routeur indécis → agent RAG par défaut. « Je ne sais pas router » ≠ « je ne réponds pas »."),
            ("dataclasses.replace —", "ajoute interaction_id à un Conseil frozen par copie (pas de mutation)."),
            ("Renvoie l'entité Conseil existante —", "tout l'aval V2 (router HTTP, DTO, disclaimer, streaming) marche sans changement. C'est ce qui permet le flag agents_enabled."),
        ],
        mental="L'orchestrateur est un routeur + garde + journaliseur. La boucle « décider → agir → vérifier » se généralise en cycles (plan-act-observe) dans les systèmes avancés ; notre version plate a la même structure.",
        code_blocks=[
            (
                "Les 7 étapes (extrait condensé) :",
                """
fil = fil_ancre(question, historique)
refus = guardrails.evaluer(fil)              # 2. garde-fou ENTRÉE
if refus is not None:
    return Conseil(refus.message, ELEVEE, [], redirection_anader=True)
agent = await self._routeur.meilleur(requete) or self._agent_de_repli()  # 3
if await self._cache.hit_rate_limit(client_ip):  # 4. rate-limit
    raise RateLimitDepasse
reponse = await agent.traiter(requete)       # 5. dispatch
if guardrails.verifier_reponse(reponse.texte):   # 6. garde-fou SORTIE
    return Conseil(guardrails.REFUS_PHYTO, ELEVEE, [], redirection_anader=True)
return await self._journaliser(question, langue, conseil)  # 7
""",
            ),
        ],
        tests=[
            "Route vers l'agent le plus pertinent (météo > rag) ; l'autre agent n'est pas appelé.",
            "Repli sur l'agent par défaut quand aucun routage n'aboutit.",
            "Une question hors filière (maïs) déclenche un refus ANADER SANS appeler d'agent.",
            "L'interaction est journalisée (interaction_id rattaché).",
            "Le rate-limit lève RateLimitDepasse avant l'inférence.",
        ],
    )

    chapter(
        doc,
        5,
        "Le squelette d'agent — services/agents/base.py + agent_rag.py",
        concept=[
            "Agentifier une capacité existante = l'envelopper dans AgentPort. Avant d'écrire "
            "4 agents qui font tous « appeler le LLM → extraire les sources → estimer la "
            "confiance → signer », on factorise cette mécanique dans AgentBase. C'est le "
            "pattern Template Method : la base définit le squelette, chaque agent ne fournit "
            "que sa spécificité (quel contexte injecter, comment scorer).",
        ],
        decisions=[
            ("AgentBase est optionnelle —", "le contrat reste un Protocol ; la base est un confort (DRY). On sépare ce qu'on DOIT respecter (contrat) de ce qu'on PEUT réutiliser (commodité)."),
            ("RAG = agent par défaut —", "généraliste ancré sur sources officielles → toujours un bon repli. Son peut_traiter renvoie un plancher modéré (0.4) : éligible partout, facile à battre par un spécialiste."),
        ],
        mental="Un agent concret = le contexte qu'il sait fabriquer + le score qu'il s'attribue. Le reste est mutualisé.",
        code_blocks=[
            (
                "AgentBase._generer (mutualisé) :",
                """
async def _generer(self, requete, contexte: str | None) -> AgentReponse:
    texte = await self._inference.generer(
        requete.question, contexte=contexte, historique=requete.historique)
    sources = postprocess.extraire_sources(texte)
    return AgentReponse(texte, sources,
        postprocess.estimer_confiance(sources), agent=self.nom)
""",
            ),
        ],
        tests=[
            "L'agent RAG passe le contexte documentaire récupéré à l'inférence.",
            "Sans récupérateur RAG, l'agent fonctionne quand même (contexte None).",
            "Son score plancher le rend éligible comme repli.",
        ],
    )

    chapter(
        doc,
        6,
        "Le tool use — services/outils/meteo.py + agent_meteo.py",
        concept=[
            "Le pattern fondateur de l'agentique. Un chatbot parle (depuis sa mémoire figée) ; "
            "un agent agit : il appelle des outils qui ramènent des données fraîches, puis "
            "raisonne dessus. L'agent Météo récupère les prévisions via OutilMeteo, les "
            "injecte comme contexte factuel dans le prompt, et le LLM raisonne sur des faits, "
            "pas sur sa mémoire (grounding). C'est le « function calling » des grands "
            "frameworks, mais explicite et déterministe.",
        ],
        decisions=[
            ("Séparer l'OUTIL de l'AGENT —", "l'outil récupère la donnée (I/O, mockable, réutilisable) ; l'agent raisonne dessus (logique, sans réseau direct). On teste chacun isolément."),
            ("Port mockable (MeteoPort) —", "aucun appel réseau en test ; la source est interchangeable. Aucun LLM tiers — données factuelles uniquement (souveraineté)."),
            ("Fail-soft —", "si l'API plante, l'outil renvoie {} au lieu d'exploser ; l'agent dégrade en conseil générique. Un outil qui échoue ne fait jamais tomber l'agent."),
            ("Routage par mots-clés —", "peut_traiter monte avec le nombre de termes climat détectés (0.7 + 0.1 × touches, plafonné à 1.0)."),
        ],
        mental="L'outil = les yeux et les mains de l'agent sur le monde réel. L'agent = le cerveau qui décide quoi en faire.",
        code_blocks=[
            (
                "L'agent récupère puis injecte (extrait) :",
                """
async def traiter(self, requete) -> AgentReponse:
    localite = _detecter_localite(requete.fil_ancre) or self._geo_defaut
    previsions = await self._outil.invoquer(localite=localite)  # tool use
    contexte = _formater_previsions(localite, previsions)
    return await self._generer(requete, contexte)  # le LLM raisonne sur les faits
""",
            ),
        ],
        tests=[
            "Score élevé sur une question météo, quasi nul sur une question prix.",
            "Les prévisions récupérées sont bien injectées dans le contexte passé au LLM.",
        ],
    )

    chapter(
        doc,
        7,
        "Réplication du pattern — services/outils/prix.py + agent_prix.py",
        concept=[
            "L'agent Prix est le jumeau de l'agent Météo : même moule (outil + port mockable "
            "+ injection de contexte), domaine différent (prix/marché du cacao). Sa valeur "
            "pédagogique : prouver que le framework tient.",
        ],
        decisions=[
            ("Coût marginal faible et constant —", "le test d'un bon socle : ajouter l'agent n°5..n°11 = recopier le moule en changeant le domaine."),
            ("invoquer(**kwargs) absorbe la différence —", "OutilMeteo prend localite, OutilPrix ne prend rien (prix national). Le contrat choisi en Task 1 paie ici."),
        ],
        mental="Quand un 2e agent « tool use » se construit en recopiant le 1er, le framework est prouvé. Le reste de la roadmap (10+ agents) est mécanique.",
        tests=[
            "Score élevé sur une question prix, quasi nul sur une question de taille.",
            "Le cours récupéré (1800 FCFA/kg) est injecté dans le contexte.",
        ],
    )

    chapter(
        doc,
        8,
        "Synthèse multi-agents — services/agents/agent_reporting.py",
        concept=[
            "Jusqu'ici un seul agent répond. L'agent Reporting est différent : il compose la "
            "sortie de PLUSIEURS agents (RAG + Météo + Prix) en une synthèse narrative. C'est "
            "le premier pas vers l'orchestration multi-agents (plusieurs contributeurs), par "
            "opposition au simple routage vers UN agent. Germe des architectures « agent "
            "superviseur ».",
        ],
        decisions=[
            ("Construit en dernier —", "il dépend des autres : il illustre qu'un agent peut consommer le travail d'agents pairs."),
            ("Agrégation prudente —", "les sources des contributions sont unionnées sans doublon ; la confiance retenue est la plus basse (prudence)."),
            ("Fusion séquentielle simple —", "la généralisation fan-out/fan-in (exécution parallèle pilotée par l'orchestrateur) est une évolution V3+ explicitement hors socle."),
        ],
        mental="Le routage choisit QUI parle ; la synthèse fait PARLER ENSEMBLE. C'est la bascule du mono-agent vers le multi-agents.",
        code_blocks=[
            (
                "synthetiser (extrait) :",
                """
async def synthetiser(self, requete, contributions: list[AgentReponse]) -> AgentReponse:
    contexte = _formater_contributions(contributions)   # « [meteo] … / [prix] … »
    base = await self._generer(requete, contexte)        # le LLM rédige la synthèse
    return AgentReponse(base.texte, _agréger_sources(contributions),
        _confiance_min(contributions) or base.confiance, agent="reporting")
""",
            ),
        ],
        tests=[
            "peut_traiter élevé sur une demande de bilan/synthèse.",
            "Les contributions sont passées au LLM comme contexte de synthèse.",
            "Les sources des contributions sont agrégées dans la réponse finale.",
        ],
    )

    # --- Câblage ---
    heading(doc, "9. Câblage derrière un flag — config, api_deps, router", level=1)
    heading(doc, "Le concept", level=3)
    para(
        doc,
        "Une plateforme agentique se met en service progressivement. On l'expose derrière "
        "agents_enabled (OFF par défaut) : tant qu'il est OFF, la V2 reste seule en "
        "production. La composition racine (api_deps) est le seul endroit où l'on ASSEMBLE "
        "le graphe d'objets : créer le registre, instancier chaque agent avec ses ports, "
        "les enregistrer, construire le routeur puis l'orchestrateur.",
    )
    heading(doc, "Les décisions de conception", level=3)
    bullets(
        doc,
        [
            ("Feature flag —", "bascule V2↔V3 sans risque ; rollback instantané."),
            ("Composition root —", "tout le câblage en un lieu ; le reste du code n'en sait rien."),
            ("Branchement via get_dialogue_service —", "le router POST passe par la gestion de sessions V2, PAS directement par le handler HTTP : on insère l'orchestrateur à ce niveau pour conserver les sessions conversationnelles."),
            ("Outils « indisponibles » —", "Météo/Prix sont enregistrés avec un adaptateur neutre ({}) tant qu'aucune API réelle n'est branchée ; l'agent dégrade en conseil générique. Le socle reste 100 % testable et déployable sans dépendance externe."),
        ],
    )
    quote(
        doc,
        "Livrer sans casser : flag OFF + composition root unique. La V3 s'insère dans la V2, elle ne la remplace pas.",
    )
    doc.add_page_break()

    # --- Recette ---
    heading(doc, "Recette — Ajouter un agent en 4 étapes", level=1)
    para(
        doc,
        "C'est l'aboutissement du socle : l'extensibilité prouvée. Pour ajouter l'agent "
        "n°5 (ex. Maladie, Satellite, Réglementation…) :",
    )
    numbered(
        doc,
        [
            ("Écrire l'agent —", "services/agents/agent_maladie.py héritant d'AgentBase, avec nom, description, mots_cles, peut_traiter(), traiter()."),
            ("(Si besoin) un outil —", "services/outils/maladie.py + un MaladiePort mockable, sur le moule de meteo.py/prix.py."),
            ("L'enregistrer —", "une ligne dans _construire_orchestrateur (api_deps) : registre.enregistrer(AgentMaladie(...))."),
            ("Tester —", "api/tests/agents/test_agent_maladie.py (routage + injection de contexte), en TDD."),
        ],
    )
    para(
        doc,
        "AUCUNE autre modification. L'orchestrateur, le registre et le routeur restent "
        "intacts. C'est la définition opérationnelle d'« ouvert à l'extension, fermé à la "
        "modification ».",
        italic=True,
        color=GREY,
    )
    doc.add_page_break()

    # --- Méthodologie ---
    heading(doc, "Méthodologie — TDD et qualité", level=1)
    bullets(
        doc,
        [
            ("Cycle TDD —", "test rouge (échoue car le code n'existe pas) → code minimal → test vert → lint → commit. Un commit par tâche."),
            ("Pourquoi le rouge d'abord —", "un test qu'on n'a jamais vu échouer ne prouve rien. Le rouge garantit qu'il teste vraiment quelque chose."),
            ("Inférence et réseau mockés —", "aucun appel réel en CI. Les ports (InferencePort, MeteoPort, PrixPort, JournalPort, CachePort) sont remplacés par des doubles."),
            ("Couverture ≥ 97 % —", "seuil CI (--cov-fail-under). Le socle a porté la suite à 508 tests verts."),
            ("Lint ruff —", "format + check ; les imports triés ; aucune exception trop large non justifiée."),
        ],
    )
    doc.add_page_break()

    # --- Garde-fous ---
    heading(doc, "Garde-fous & souveraineté (non négociable)", level=1)
    bullets(
        doc,
        [
            ("Périmètre cacao UNIQUEMENT —", "vivrier, anacarde, médical, dosages phytosanitaires précis → redirection ANADER. Décision Waopron, juin 2026."),
            ("Garde-fous dans l'orchestrateur —", "centralisés, jamais réimplémentés par agent. S'appliquent identiquement à tous les agents, actuels et futurs."),
            ("Aucun service externe en production —", "OpenAI/Anthropic/Cohere proscrits dans le pipeline. Les outils appellent des sources de données, pas des LLM tiers, toujours derrière un port mockable."),
            ("Disclaimer ANADER systématique —", "porté par l'entité Conseil sur chaque réponse modèle."),
            ("Jamais de dosage généré —", "même en exemple de test."),
        ],
    )
    doc.add_page_break()

    # --- Glossaire ---
    heading(doc, "Glossaire de l'IA agentique", level=1)
    table(
        doc,
        ["Terme", "Définition"],
        [
            ["Agent", "Capacité bornée derrière une interface stable (peut_traiter + traiter)."],
            ["Outil (tool)", "Fonction nommée à entrée/sortie sérialisables qu'un agent invoque pour agir."],
            ["Tool use", "Mécanisme par lequel un agent appelle un outil pour ramener des données fraîches puis raisonne dessus."],
            ["Grounding", "Ancrage de la génération sur des faits injectés en contexte plutôt que sur la mémoire du modèle."],
            ["Routage d'intention", "Choix du/des agent(s) pertinent(s) pour une requête (ici par score déterministe)."],
            ["Orchestrateur", "Plan de contrôle : décide qui agit, applique garde-fous, dispatche, journalise."],
            ["Registre", "Annuaire d'agents ; point d'extension du framework."],
            ["Composition root", "Lieu unique où le graphe d'objets est assemblé (câblage des dépendances)."],
            ["Feature flag", "Interrupteur de configuration activant une capacité sans redéploiement de code."],
            ["Fail-soft", "Dégradation propre : un composant qui échoue ne fait pas tomber l'ensemble."],
            ["Défense en profondeur", "Plusieurs contrôles indépendants (entrée + sortie) plutôt qu'un seul."],
            ["Template Method", "Pattern où une base définit le squelette d'un algorithme, les sous-classes la spécificité."],
            ["Plan-act-observe", "Boucle agentique multi-étapes (planifier, agir, observer) — généralisation de l'orchestrateur plat."],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    build()
