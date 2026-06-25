"""Construit le business plan investisseurs d'OpenLab Consulting (Word).

Portefeuille de projets R&D en IA — Côte d'Ivoire d'abord. Page de garde avec
logo, résumé exécutif, contexte marché chiffré et sourcé, un chapitre approfondi
par projet (problème, solution, marché, architecture, réussite/échec, modèle
économique), synthèse comparative, recommandation et annexe de sources.

Les visuels proviennent de scripts/gen_bp_assets.py (docs/img_bp/). Les chiffres
sont ceux, sourcés, de gen_bp_assets.DATA.

Usage : python scripts/gen_bp_assets.py && python scripts/build_businessplan.py
Sortie : docs/OpenLab_Consulting_Business_Plan.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from gen_bp_assets import DATA, SOURCES

ROOT = Path(__file__).resolve().parents[1]
IMG = ROOT / "docs" / "img_bp"
LOGO = ROOT / "docs" / "OPENLAB.png"
OUT = ROOT / "docs" / "OpenLab_Consulting_Business_Plan.docx"

OR = RGBColor(0xEA, 0x5B, 0x13)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GREY = RGBColor(0x60, 0x60, 0x60)

ACC = {"OR": "EA5B13", "GREEN": "2E7D32", "BLUE": "1565C0", "PURPLE": "6A1B9A", "TEAL": "00838F"}


def _set_font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def heading(doc, text, level=1, color=OR):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    r.font.color.rgb = color
    return p


def para(doc, text, *, size=11, italic=False, color=DARK, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    _set_font(p.add_run(text), size=size, italic=italic, color=color)
    return p


def lead(doc, label, text):
    """Paragraphe « Intitulé. texte » avec l'intitulé en gras coloré."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _set_font(p.add_run(label + " "), bold=True, color=OR)
    _set_font(p.add_run(text))
    return p


def bullets(doc, items, color=DARK):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_font(p.add_run(it), color=color)


def image(doc, name, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(IMG / name), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    _set_font(c.add_run(caption), size=9, italic=True, color=GREY)


def stat_table(doc, theme, rows=6):
    data = DATA[theme][:rows]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(("Indicateur", "Valeur", "Source (année)")):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for libelle, val, an, src, _url, _conf in data:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = libelle, val, f"{src} ({an})"
        for j in range(3):
            for pr in c[j].paragraphs:
                for rn in pr.runs:
                    rn.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break(doc):
    doc.add_page_break()


def hrule(p, color="EA5B13"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def footer(doc):
    sec = doc.sections[0]
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("OpenLab Consulting — Confidentiel · "), size=8, color=GREY)
    # champ PAGE
    run = p.add_run()
    fld1, instr, fld2 = OxmlElement("w:fldChar"), OxmlElement("w:instrText"), OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    _set_font(run, size=8, color=GREY)


# --------------------------------------------------------------------------- #
# Contenu des projets (Côte d'Ivoire d'abord)                                 #
# --------------------------------------------------------------------------- #
PROJETS = [
    {
        "nom": "Data Factory CI", "score": 95, "acc": "OR", "theme": "ia_data",
        "chart": "ai_africa.png", "chart_cap": "Marché de l'IA en Afrique : 4,5 → 16,5 Md$ d'ici 2030 (TCAC 27,4 %).",
        "archi": "archi_datafactory.png",
        "soustitre": "Construire les fondations Data/IA souveraines des entreprises et institutions ivoiriennes.",
        "probleme": "Banques, télécoms, agro-industries et administrations ivoiriennes accumulent des données sans "
                    "infrastructures fiables, gouvernées et souveraines pour les exploiter. Les compétences clés "
                    "(gouvernance, MLOps, plateformes analytiques) sont rares et majoritairement importées à coût élevé.",
        "solution": "OpenLab conçoit et opère une « usine à données » clé en main : ingestion/ETL, data lake gouverné, "
                    "MLOps, plateformes analytiques et déploiement souverain. Le socle est réutilisable d'un client à "
                    "l'autre, ce qui industrialise la livraison et améliore les marges.",
        "ci": "La Côte d'Ivoire est le moteur économique de l'UEMOA (PIB 86,9 Md$, croissance 6 %/an) avec un tissu "
              "d'entreprises et d'institutions en pleine numérisation, alors que l'offre Data/IA locale est quasi "
              "inexistante face à un marché IA africain en forte croissance.",
        "succes": [
            "Capacité bout-en-bout déjà prouvée par OpenCacao (LLM souverain, MLOps, déploiement).",
            "Revenus récurrents : exploitation, maintenance et MLOps en plus du projet initial.",
            "Barrière à l'entrée élevée : rareté des compétences + exigence de souveraineté.",
            "Effet plateforme : un socle réutilisable réduit le coût marginal de chaque nouveau client.",
            "Ancrage local : alternative crédible et moins coûteuse aux cabinets et hyperscalers étrangers.",
        ],
        "risques": [
            "Cycle de vente B2B long auprès des grands comptes.",
            "Pénurie de talents data/IA — à mitiger par une académie interne de formation.",
            "Dépendance initiale à quelques comptes de référence.",
            "Concurrence des hyperscalers et grandes SSII sur les gros marchés.",
            "Investissement d'infrastructure et d'avant-vente initial à financer.",
        ],
        "modele": "« Build » (projet d'intégration) + « Run » (abonnement exploitation/MLOps) + formation. "
                  "Cibles : banques, télécoms, agro-industrie, institutions publiques et bailleurs.",
    },
    {
        "nom": "IA Assurance Agricole", "score": 85, "acc": "GREEN", "theme": "assurance",
        "chart": "assurance_penetration.png", "chart_cap": "~99 % des petits exploitants africains ne sont pas assurés : un marché quasi vierge.",
        "archi": "archi_assurance.png",
        "soustitre": "Scoring de risque et assurance indicielle pour banques, assurances et microfinances.",
        "probleme": "L'agriculture représente ~18 % du PIB et ~45 % de l'emploi en Côte d'Ivoire, mais près de 99 % des "
                    "petits exploitants africains ne sont pas assurés. Les assureurs manquent de données fiables pour "
                    "tarifer un risque climatique croissant, et le crédit agricole reste bridé par l'absence de garantie.",
        "solution": "Une plateforme de scoring du risque agricole et d'assurance indicielle (paramétrique) fondée sur "
                    "données satellitaires, météo et rendements. Une API alimente assureurs, banques et microfinances ; "
                    "l'indemnisation se déclenche automatiquement sur un indice objectif, sans expertise terrain coûteuse.",
        "ci": "La Côte d'Ivoire combine une agriculture massive (cacao en tête), une exposition climatique croissante et "
              "un écosystème mobile money (> 5 % du PIB) idéal pour distribuer les polices et verser les indemnités.",
        "succes": [
            "Marché mondial de l'assurance agricole de 41,5 à 70 Md$ ; pénétration africaine ~1 % = potentiel immense.",
            "L'indice paramétrique élimine l'essentiel des coûts d'expertise et de la fraude.",
            "Synergie directe avec OpenCacao : données et confiance déjà établies dans la filière.",
            "Effet de levier sur le crédit agricole (dé-risquage) attractif pour banques et bailleurs.",
            "Distribution via mobile money et coopératives déjà numérisées.",
        ],
        "risques": [
            "Qualité et granularité des données satellitaires/météo.",
            "Risque de base (écart entre indice et perte réelle) à calibrer finement.",
            "Cadre réglementaire assurantiel régional (CIMA) à respecter.",
            "Confiance et éducation des assurés à construire.",
            "Sinistralité corrélée (sécheresse systémique) imposant une réassurance.",
        ],
        "modele": "Scoring/SaaS à l'usage pour assureurs et banques + commission sur primes + partenariats microfinance. "
                  "Pilote ciblé sur une zone cacao avec un assureur et une institution de microfinance.",
    },
    {
        "nom": "OpenCompliance AI", "score": 80, "acc": "BLUE", "theme": "compliance",
        "chart": "regtech.png", "chart_cap": "RegTech mondiale : 20,7 → 44,1 Md$ d'ici 2030 (TCAC 16,4 %).",
        "archi": "archi_compliance.png",
        "soustitre": "Assistant de conformité réglementaire souverain (RAG + IA générative) pour la place financière UEMOA.",
        "probleme": "Banques, assurances, fintechs et grandes entreprises subissent une pression réglementaire croissante "
                    "(LBC/FT, KYC, normes BCEAO/UEMOA, protection des données). La conformité est coûteuse, manuelle et "
                    "exposée à de lourdes amendes.",
        "solution": "Un assistant de conformité ancré sur les textes (RAG) qui répond avec citations vérifiables, exécute "
                    "des contrôles et workflows, génère des dossiers et tient un journal d'audit — le tout souverain, en "
                    "français et adapté au cadre régional.",
        "ci": "Abidjan est la place financière de l'UEMOA, avec un essor fintech et mobile money sous une réglementation "
              "BCEAO exigeante : un besoin direct d'outils de conformité localisés et de confiance.",
        "succes": [
            "RegTech à 16,4 % de TCAC ; coût de la conformité de 61 Md$ (US) à 85 Md$ (EMEA) par an.",
            "Amendes AML/KYC de 4,6 Md$ en 2024 : un ROI de prévention immédiat.",
            "Réponses citées et traçables : la confiance, condition d'adoption en conformité.",
            "Souveraineté et résidence des données : atout décisif vs solutions étrangères.",
            "Réutilisation directe de la pile RAG/LLM d'OpenCacao.",
        ],
        "risques": [
            "Exigence de fiabilité quasi absolue (responsabilité en cas d'erreur).",
            "Accès à des textes réglementaires à jour et structurés.",
            "Cycle de vente régulé et long dans le secteur financier.",
            "Question de responsabilité juridique du conseil automatisé.",
            "Concurrence de RegTech internationales mieux financées.",
        ],
        "modele": "SaaS par siège + module audit + intégration. Cibles : banques, assurances, fintechs et grandes "
                  "entreprises soumises à la conformité régionale.",
    },
    {
        "nom": "IA Marchés Publics", "score": 75, "acc": "PURPLE", "theme": "procurement",
        "chart": "procurement_gdp.png", "chart_cap": "La commande publique pèse ~17 % du PIB en Afrique (13 % OCDE).",
        "chart2": "procurement_waste.png", "chart2_cap": "~25 % de la commande publique mondiale est gaspillée (Banque mondiale).",
        "archi": "archi_marches.png",
        "soustitre": "Analyse automatique des appels d'offres et assistance à la réponse — pour les PME et les acheteurs publics.",
        "probleme": "La commande publique représente ~17 % du PIB en Afrique, dont près d'un quart se perd en gaspillage "
                    "et 10 à 25 % par contrat peut être détourné. Côté entreprises, répondre aux appels d'offres reste "
                    "complexe, chronophage et inaccessible à beaucoup de PME.",
        "solution": "Une plateforme qui agrège les appels d'offres, en extrait l'information par NLP, détecte les "
                    "opportunités pertinentes, assiste la rédaction des réponses et alerte sur les échéances. Côté "
                    "acheteurs, elle outille l'analyse et la transparence.",
        "ci": "La Côte d'Ivoire conjugue une commande publique importante, un agenda de transparence et de numérisation, "
              "et un tissu de PME mal outillées pour soumissionner.",
        "succes": [
            "Marché des logiciels d'achat de 10,7 à 17,1 Md$ ; valeur double (PME soumissionnaires + acheteurs).",
            "Gains anti-gaspillage et anti-corruption mesurables, alignés sur les priorités des bailleurs.",
            "Concurrence locale faible ; barrière liée à la donnée et à l'expertise métier.",
            "Réutilisation du moteur NLP/LLM souverain d'OpenLab.",
            "Effet réseau : plus de PME et d'AO référencés, plus la plateforme est utile.",
        ],
        "risques": [
            "Dépendance à l'ouverture et à la qualité des données publiques.",
            "Sensibilité politique du sujet transparence/anti-corruption.",
            "Fragmentation des portails et formats d'appels d'offres.",
            "Cycle d'achat public long côté acheteurs institutionnels.",
            "Adoption des PME à accompagner (formation, confiance).",
        ],
        "modele": "SaaS PME (abonnement + alertes) + licences institutions/acheteurs + facturation à l'usage (réponses "
                  "assistées). Démarrage par l'agrégation des AO ivoiriens et un groupement de PME pilote.",
    },
    {
        "nom": "OpenGov AI", "score": 70, "acc": "TEAL", "theme": "govtech",
        "chart": "govtech.png", "chart_cap": "GovTech mondiale : 825 → 3 091 Md$ d'ici 2035 (TCAC 15,8 %).",
        "chart2": "egdi.png", "chart2_cap": "L'Afrique (EGDI 0,42) reste sous la moyenne mondiale (0,64) : marge de progression.",
        "archi": "archi_gov.png",
        "soustitre": "Assistant intelligent multicanal pour les administrations et les usagers des services publics.",
        "probleme": "L'administration ivoirienne se numérise mais l'accès aux services publics reste complexe ; l'indice "
                    "africain d'e-gouvernement (0,42) est sous la moyenne mondiale (0,64). Les agents sont surchargés et "
                    "l'information dispersée.",
        "solution": "Un assistant multicanal (web, WhatsApp, USSD) pour citoyens et agents, ancré sur les procédures "
                    "officielles (RAG), intégré au back-office, avec des analytics de pilotage — en français et langues "
                    "locales, en mode souverain.",
        "ci": "Il existe une volonté politique d'e-gouvernement, un taux d'équipement mobile très élevé (149 %) et un "
              "besoin d'accès en langues locales ; la souveraineté est une exigence forte pour un acheteur public.",
        "succes": [
            "GovTech à 15,8 % de TCAC ; l'écart d'EGDI mesure directement la marge de progression.",
            "Impact social et d'inclusion fort (accès équitable aux services).",
            "Relation institutionnelle structurante et effet vitrine national.",
            "Souveraineté et résidence des données : critère décisif pour l'État.",
            "Réutilisation de l'assistant conversationnel multicanal déjà éprouvé (OpenCacao V2/V3).",
        ],
        "risques": [
            "Cycle de vente public très long et fortement politique.",
            "Exigences de sécurité et de souveraineté élevées.",
            "Dépendance aux budgets et aux priorités publics.",
            "Conduite du changement auprès des agents.",
            "Responsabilité sur l'exactitude de l'information officielle.",
        ],
        "modele": "Licence/abonnement institutionnel + intégration + maintenance, en déploiement souverain. Pilote sur "
                  "un service public à fort volume de demandes.",
    },
]


def build():
    doc = Document()
    # Marges un peu resserrées pour les visuels larges.
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
    footer(doc)

    # ----------------------- PAGE DE GARDE ----------------------- #
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(str(LOGO), width=Inches(2.6))
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(t.add_run("Portefeuille de projets R&D en Intelligence Artificielle"), size=24, bold=True, color=DARK)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(st.add_run("Cinq opportunités à fort potentiel pour la Côte d'Ivoire"), size=14, italic=True, color=OR)
    rule_p = doc.add_paragraph()
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrule(rule_p)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    meta = [
        ("Éditeur", "OpenLab Consulting — projets de R&D"),
        ("Objet", "Business plan — portefeuille IA souveraine, Côte d'Ivoire d'abord"),
        ("Auteur / Direction", "Waopron Coulibaly"),
        ("Document", "Confidentiel — destiné aux investisseurs et partenaires"),
        ("Version", "1.0 — 25 juin 2026"),
    ]
    mt = doc.add_table(rows=0, cols=2)
    mt.alignment = 1
    for k, v in meta:
        c = mt.add_row().cells
        rk = c[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.color.rgb = OR
        rk.font.size = Pt(10)
        rv = c[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(sig.add_run("OpenLab Consulting — IA souveraine pour la Côte d'Ivoire"), size=11, bold=True, color=OR)
    page_break(doc)

    # ----------------------- 1. RÉSUMÉ EXÉCUTIF ----------------------- #
    heading(doc, "1. Résumé exécutif", 1)
    para(doc, "OpenLab Consulting est un acteur ivoirien de la Data et de l'Intelligence Artificielle. Sa thèse : "
              "bâtir, depuis la Côte d'Ivoire, des solutions d'IA souveraines — données résidentes, modèles ouverts "
              "auto-hébergés, sans dépendance à des services tiers en production. Le démonstrateur OpenCacao (assistant "
              "agronomique cacao, LLM affiné servi en propre) prouve cette capacité de bout en bout.")
    para(doc, "Ce document présente cinq projets R&D priorisés pour le marché ivoirien, puis ouest-africain. Ils "
              "partagent une même fabrique technologique (LLM souverain, RAG, MLOps) et un même avantage défendable : "
              "la maîtrise de la donnée métier locale et de la souveraineté, là où l'offre internationale est coûteuse "
              "et mal adaptée.")
    image(doc, "portfolio.png", "Figure 1 — Probabilité de réussite des cinq projets du portefeuille (évaluation OpenLab Consulting).", width=5.6)
    para(doc, "Le contexte est porteur : un marché de l'IA africain qui passerait de 4,5 à 16,5 Md$ d'ici 2030 "
              "(TCAC 27 %), une économie ivoirienne en croissance de 6 %/an et un taux d'équipement mobile de 149 %. "
              "La Côte d'Ivoire, premier producteur mondial de cacao (44 % de l'offre), offre en outre des verticales "
              "métier à fort impact (agriculture, finance, secteur public).")
    page_break(doc)

    # ----------------------- 2. CONTEXTE CÔTE D'IVOIRE ----------------------- #
    heading(doc, "2. Contexte : la Côte d'Ivoire d'abord", 1)
    para(doc, "Le portefeuille vise en priorité le marché ivoirien — base de référence, de données et de premières "
              "ventes — avant un déploiement régional (UEMOA/CEDEAO) puis continental. Les indicateurs ci-dessous "
              "cadrent la taille et la maturité du marché adressable.")
    heading(doc, "2.1  Indicateurs macro & numériques", 2)
    stat_table(doc, "cote_ivoire")
    heading(doc, "2.2  Une vague IA continentale", 2)
    image(doc, "ai_africa.png", "Figure 2 — Marché de l'IA en Afrique : 4,5 → 16,5 Md$ d'ici 2030 (Mastercard / Statista).", width=5.4)
    heading(doc, "2.3  Méthode d'évaluation du portefeuille", 2)
    para(doc, "Chaque projet est noté selon quatre critères : (1) taille et croissance du marché, (2) barrière à "
              "l'entrée et caractère défendable, (3) potentiel de revenus récurrents, (4) adéquation avec la "
              "souveraineté et la donnée métier déjà maîtrisée. La probabilité de réussite synthétise ces critères.")
    page_break(doc)

    # ----------------------- 3. PROJETS ----------------------- #
    heading(doc, "3. Les cinq projets en profondeur", 1)
    for i, pr in enumerate(PROJETS, start=1):
        acc = RGBColor.from_string(ACC[pr["acc"]])
        heading(doc, f"3.{i}  {pr['nom']}  —  {pr['score']} %", 2, color=acc)
        para(doc, pr["soustitre"], italic=True, color=GREY)
        lead(doc, "Le problème.", pr["probleme"])
        lead(doc, "La solution.", pr["solution"])
        lead(doc, "Pourquoi la Côte d'Ivoire d'abord.", pr["ci"])
        heading(doc, "Marché (chiffres sourcés)", 3, color=acc)
        stat_table(doc, pr["theme"])
        image(doc, pr["chart"], f"Figure 3.{i} — {pr['chart_cap']}", width=5.4)
        if pr.get("chart2"):
            image(doc, pr["chart2"], f"Figure 3.{i}b — {pr['chart2_cap']}", width=5.0)
        heading(doc, "Architecture cible", 3, color=acc)
        image(doc, pr["archi"], f"Figure 3.{i}c — Architecture souveraine de {pr['nom']}.", width=6.3)
        heading(doc, "Pourquoi cela peut réussir considérablement", 3, color=acc)
        bullets(doc, pr["succes"])
        heading(doc, "Ce qui pourrait la faire échouer (et nos parades)", 3, color=acc)
        bullets(doc, pr["risques"])
        lead(doc, "Modèle économique & go-to-market.", pr["modele"])
        page_break(doc)

    # ----------------------- 4. SYNTHÈSE ----------------------- #
    heading(doc, "4. Synthèse comparative & recommandation", 1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("Projet", "Réussite", "Marché de référence", "Logique de revenus")):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    synth = [
        ("Data Factory CI", "95 %", "IA Afrique 4,5→16,5 Md$ (27 % TCAC)", "Build + Run (MLOps) récurrent"),
        ("IA Assurance Agricole", "85 %", "Assurance agricole 41,5→70 Md$ ; ~1 % d'assurés", "Scoring à l'usage + commissions"),
        ("OpenCompliance AI", "80 %", "RegTech 20,7→44,1 Md$ (16 % TCAC)", "SaaS par siège + audit"),
        ("IA Marchés Publics", "75 %", "Logiciels d'achat 10,7→17,1 Md$", "SaaS PME + licences institutions"),
        ("OpenGov AI", "70 %", "GovTech 825→3 091 Md$ (16 % TCAC)", "Licence institutionnelle + intégration"),
    ]
    for nom, sc, march, rev in synth:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = nom, sc, march, rev
        for j in range(4):
            for p in c[j].paragraphs:
                for rn in p.runs:
                    rn.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    heading(doc, "Recommandation", 2)
    para(doc, "Deux moteurs jumeaux portent le portefeuille : Data Factory CI (la fabrique d'infrastructures, qui "
              "industrialise tous les autres projets) et la suite OpenCacao (l'avantage métier déjà acquis, traité dans "
              "son propre backlog V3). Sur ce socle, IA Assurance Agricole et IA Marchés Publics sont des verticales à "
              "fort impact et à concurrence locale faible. OpenGov AI constitue le projet vitrine institutionnel, à "
              "cycle plus long mais à effet de réputation majeur.")
    para(doc, "Séquencement R&D conseillé : (1) consolider Data Factory CI et la suite OpenCacao comme socles ; "
              "(2) lancer en parallèle un pilote Assurance Agricole (synergie cacao) et Marchés Publics (faible "
              "barrière, valeur rapide) ; (3) engager OpenGov AI via un partenariat institutionnel. Chaque projet "
              "réutilise la même pile souveraine, ce qui mutualise l'investissement et accélère le temps de marché.")
    page_break(doc)

    # ----------------------- ANNEXE SOURCES ----------------------- #
    heading(doc, "Annexe — Sources", 1)
    para(doc, "Toutes les statistiques de ce document sont attribuées à une source nommée et datée. Liste complète :",
         italic=True, color=GREY)
    for libelle, val, an, src, url in SOURCES:
        p = doc.add_paragraph(style="List Bullet")
        _set_font(p.add_run(f"{libelle} : "), bold=True)
        _set_font(p.add_run(f"{val} ({an}) — {src}. "))
        _set_font(p.add_run(url), size=8, color=RGBColor(0x15, 0x65, 0xC0))

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    build()
